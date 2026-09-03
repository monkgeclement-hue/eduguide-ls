import json
import hashlib
import smtplib
import mimetypes
import os
import re
import sqlite3
import secrets
import tempfile
import uuid
from datetime import datetime, timedelta, timezone
from email.message import EmailMessage
from pathlib import Path
from typing import Any
from urllib import error as urlerror
from urllib import request as urlrequest
from urllib.parse import quote

from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, Header, HTTPException, Request, UploadFile
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.responses import FileResponse, Response
from openai import OpenAI
from pydantic import BaseModel, Field


ROOT = Path(__file__).resolve().parent
load_dotenv(ROOT / ".env")
DB_PATH = ROOT / "data" / "eduguide.db"
UPLOAD_ROOT = ROOT / "data" / "uploads"
MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_FILES_PER_UPLOAD = 5
ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".docx", ".jpg", ".jpeg", ".png"}
DEFAULT_SUPABASE_STORAGE_BUCKET = "eduguide-documents"
AI_CHAT_HISTORY_LIMIT = 24
AI_CHAT_CONTEXT_LIMIT = 12
AI_GUIDANCE_MAX_MATCHES = 12
AI_GUIDANCE_MAX_BLOCKED_MATCHES = 8
AI_GUIDANCE_MAX_DOCUMENTS = 12
AI_GUIDANCE_MAX_QUESTION_LENGTH = 900
AI_GUIDANCE_MAX_MESSAGE_LENGTH = 1200
AI_GUIDANCE_MAX_CONVERSATION_ITEMS = 12
AI_GUIDANCE_MAX_PAYLOAD_BYTES = 150000
EMAIL_VERIFICATION_TTL_MINUTES = 10
EMAIL_VERIFICATION_RESEND_SECONDS = 45
EMAIL_VERIFICATION_MAX_ATTEMPTS = 6
AUTH_SESSION_TTL_DAYS = max(1, int(os.getenv("AUTH_SESSION_TTL_DAYS", "30") or "30"))
AUTH_SESSION_TOKEN_MIN_LENGTH = 24
AUTH_SESSION_TOKEN_MAX_LENGTH = 256
RATE_LIMIT_STATE: dict[str, list[datetime]] = {}
SECURITY_CLEANUP_INTERVAL_SECONDS = 3600
LAST_SECURITY_CLEANUP_AT: datetime | None = None
GRADE_VALUES = {"A*", "A", "B", "C", "D", "E", "F", "G", "X", "Z"}
GEMINI_DEFAULT_MODEL = "gemini-3.6-flash"
GEMINI_MODEL_FALLBACKS = ["gemini-3.6-flash", "gemini-2.5-flash", "gemini-flash-latest"]
DEPRECATED_GEMINI_MODELS = {
  "gemini-2.5-pro": GEMINI_DEFAULT_MODEL,
  "gemini-1.5-flash": GEMINI_DEFAULT_MODEL,
  "gemini-1.5-pro": GEMINI_DEFAULT_MODEL,
}
SUBJECT_ALIASES = [
  {"code": "MATH", "subject": "Mathematics", "aliases": ["mathematics", "maths", "math"]},
  {"code": "ENG", "subject": "English", "aliases": ["english language", "english"]},
  {"code": "SES", "subject": "Sesotho", "aliases": ["sesotho"]},
  {"code": "PSCI", "subject": "Physical Science", "aliases": ["physical science", "physical sciences", "double science", "double sciences"]},
  {"code": "CSK", "subject": "Computer Skills", "aliases": ["computer skills", "computer studies", "computer literacy", "computer applications", "ict skills", "information and communication technology"]},
  {"code": "ACC", "subject": "Accounting", "aliases": ["accounting", "accounts"]},
  {"code": "BIO", "subject": "Biology", "aliases": ["biology", "life science"]},
  {"code": "AGR", "subject": "Agriculture", "aliases": ["agriculture", "agricultural science"]},
  {"code": "FNU", "subject": "Food & Nutrition", "aliases": ["food and nutrition", "food & nutrition", "food nutrition", "nutrition", "food studies", "home economics", "consumer science"]},
  {"code": "REL", "subject": "Religious Knowledge", "aliases": ["religious knowledge", "religious education", "religion", "divinity", "bible knowledge"]},
  {"code": "HIST", "subject": "History", "aliases": ["history"]},
  {"code": "GEO", "subject": "Geography", "aliases": ["geography"]},
  {"code": "PHY", "subject": "Physics", "aliases": ["physics"]},
  {"code": "CHEM", "subject": "Chemistry", "aliases": ["chemistry"]},
  {"code": "ECON", "subject": "Economics", "aliases": ["economics", "economy"]},
  {"code": "LIT", "subject": "English Literature", "aliases": ["english literature", "literature in english", "literature"]},
]

app = FastAPI(title="EduGuide LS AI Server")
app.add_middleware(GZipMiddleware, minimum_size=1024)
PUBLIC_DATA_FILES = {"admin-catalog.js", "catalog.js", "source-manifest.json", "supabase-config.js"}
PUBLIC_ICON_FILES = {"icon-192.svg", "icon-512.svg", "icon-192.png", "icon-512.png", "apple-touch-icon.png"}
mimetypes.add_type("application/manifest+json", ".webmanifest")
mimetypes.add_type("image/svg+xml", ".svg")
CSP_POLICY = (
  "default-src 'self'; "
  "base-uri 'self'; "
  "object-src 'none'; "
  "frame-ancestors 'none'; "
  "form-action 'self'; "
  "connect-src 'self' https://*.supabase.co https://unpkg.com https://cdn.jsdelivr.net; "
  "img-src 'self' data: blob: https://*.supabase.co; "
  "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com https://fonts.gstatic.com; "
  "font-src 'self' data: https://fonts.gstatic.com https://fonts.googleapis.com; "
  "script-src 'self' 'unsafe-inline' https://unpkg.com https://cdn.jsdelivr.net; "
  "script-src-elem 'self' 'unsafe-inline' https://unpkg.com https://cdn.jsdelivr.net; "
  "media-src 'self' blob:; "
  "manifest-src 'self'; "
  "worker-src 'self' blob:;"
)
NO_CACHE_HEADERS = {"Cache-Control": "no-cache, max-age=0, must-revalidate", "X-Content-Type-Options": "nosniff"}
STATIC_CACHE_HEADERS = {"Cache-Control": "public, max-age=300, stale-while-revalidate=86400", "X-Content-Type-Options": "nosniff"}
IMMUTABLE_CACHE_HEADERS = {"Cache-Control": "public, max-age=604800, immutable", "X-Content-Type-Options": "nosniff"}
SECURITY_HEADERS = {
  "Strict-Transport-Security": "max-age=31536000; includeSubDomains",
  "X-Frame-Options": "DENY",
  "X-Content-Type-Options": "nosniff",
  "Referrer-Policy": "strict-origin-when-cross-origin",
  "Permissions-Policy": "camera=(), microphone=(), geolocation=(), interest-cohort=()",
  "Cross-Origin-Opener-Policy": "same-origin",
  "Cross-Origin-Resource-Policy": "same-origin",
  "Content-Security-Policy": CSP_POLICY,
}


@app.middleware("http")
async def add_security_headers(request: Request, call_next):
  response = await call_next(request)
  for header_name, header_value in SECURITY_HEADERS.items():
    response.headers.setdefault(header_name, header_value)
  return response


class GuidanceRequest(BaseModel):
  profile: dict[str, Any] = Field(default_factory=dict)
  readiness: dict[str, Any] = Field(default_factory=dict)
  matches: list[dict[str, Any]] = Field(default_factory=list)
  blockedMatches: list[dict[str, Any]] = Field(default_factory=list)
  documents: list[dict[str, Any]] = Field(default_factory=list)
  conversation: list[dict[str, Any]] = Field(default_factory=list)
  question: str | None = None
  mode: str = "guidance"


class AuthLoginRequest(BaseModel):
  email: str
  password: str


class AuthRegisterRequest(BaseModel):
  name: str
  email: str
  password: str
  district: str | None = None
  code: str | None = None


class AuthVerifyRegistrationRequest(AuthRegisterRequest):
  code: str


class AuthPasswordResetRequest(BaseModel):
  email: str


class AuthPasswordResetConfirmRequest(BaseModel):
  email: str
  code: str
  password: str


class AuthProfileRequest(BaseModel):
  name: str | None = None
  district: str | None = None
  stream: str | None = None
  leavingYear: str | None = None
  incomeBand: str | None = None
  needSignals: list[str] = Field(default_factory=list)
  preferenceText: str | None = None
  grades: dict[str, str] = Field(default_factory=dict)
  documents: list[dict[str, Any]] = Field(default_factory=list)
  shortlist: list[str] = Field(default_factory=list)
  shortlistPathways: dict[str, str] = Field(default_factory=dict)


class UserRoleUpdate(BaseModel):
  role: str
  institution: str | None = None
  managedInstitution: str | None = None


class UserStatusUpdate(BaseModel):
  status: str


class AdminTestEmailRequest(BaseModel):
  email: str | None = None


class RuntimeEventRequest(BaseModel):
  eventType: str | None = None
  event_type: str | None = None
  label: str | None = None
  payload: dict[str, Any] = Field(default_factory=dict)


class InstitutionProposalRequest(BaseModel):
  programmeId: str
  programmeName: str | None = None
  institution: str | None = None
  changes: dict[str, Any] = Field(default_factory=dict)
  note: str | None = None


class InstitutionProposalDecisionRequest(BaseModel):
  status: str
  note: str | None = None




def get_gemini_model_candidates() -> list[str]:
  requested = (os.getenv("GEMINI_MODEL") or GEMINI_DEFAULT_MODEL).strip() or GEMINI_DEFAULT_MODEL
  requested = DEPRECATED_GEMINI_MODELS.get(requested, requested)
  candidates = [requested, *GEMINI_MODEL_FALLBACKS]
  seen = set()
  return [model for model in candidates if model and not (model in seen or seen.add(model))]

def get_db_connection() -> sqlite3.Connection:
  DB_PATH.parent.mkdir(parents=True, exist_ok=True)
  connection = sqlite3.connect(DB_PATH)
  connection.row_factory = sqlite3.Row
  connection.execute("pragma busy_timeout = 5000")
  connection.execute("pragma journal_mode = wal")
  return connection


def init_database() -> None:
  with get_db_connection() as connection:
    connection.execute(
      """
      create table if not exists app_state (
        state_key text primary key,
        payload text not null,
        updated_at text not null default current_timestamp
      )
      """
    )
    connection.execute(
      """
      create table if not exists recommendation_runs (
        id integer primary key autoincrement,
        mode text not null default 'guidance',
        question text,
        profile_name text,
        payload text not null,
        created_at text not null default current_timestamp
      )
      """
    )
    connection.execute(
      """
      create table if not exists ai_chat_messages (
        id text primary key,
        user_id text not null,
        role text not null,
        content text not null,
        payload text,
        created_at text not null default current_timestamp
      )
      """
    )
    connection.execute(
      """
      create index if not exists idx_ai_chat_messages_user_created
      on ai_chat_messages(user_id, created_at desc)
      """
    )
    connection.execute(
      """
      create table if not exists uploaded_documents (
        id text primary key,
        user_id text not null,
        original_name text not null,
        stored_name text not null,
        content_type text,
        size_bytes integer not null,
        storage_path text not null,
        status text not null default 'Uploaded - OCR pending',
        extraction_status text not null default 'pending',
        extraction_text text,
        extracted_grades text not null default '[]',
        extraction_error text,
        extracted_at text,
        uploaded_at text not null default current_timestamp
      )
      """
    )
    connection.execute(
      """
      create table if not exists auth_sessions (
        token text primary key,
        user_id text not null,
        created_at text not null default current_timestamp,
        last_seen_at text not null default current_timestamp
      )
      """
    )
    connection.execute(
      """
      create table if not exists email_verifications (
        id text primary key,
        email text not null,
        purpose text not null default 'registration',
        code_hash text not null,
        code_salt text not null,
        payload text not null default '{}',
        attempts integer not null default 0,
        created_at text not null,
        expires_at text not null,
        consumed_at text
      )
      """
    )
    connection.execute(
      """
      create table if not exists runtime_events (
        id text primary key,
        user_id text not null,
        event_type text not null,
        label text,
        payload text not null default '{}',
        created_at text not null default current_timestamp
      )
      """
    )
    connection.execute(
      """
      create index if not exists idx_runtime_events_type_created
      on runtime_events(event_type, created_at desc)
      """
    )
    connection.execute(
      """
      create index if not exists idx_runtime_events_user_created
      on runtime_events(user_id, created_at desc)
      """
    )
    connection.execute(
      """
      create index if not exists idx_email_verifications_email_purpose_created
      on email_verifications(email, purpose, created_at desc)
      """
    )
    existing_columns = {
      row["name"]
      for row in connection.execute("pragma table_info(uploaded_documents)").fetchall()
    }
    document_migrations = {
      "extraction_status": "alter table uploaded_documents add column extraction_status text not null default 'pending'",
      "extraction_text": "alter table uploaded_documents add column extraction_text text",
      "extracted_grades": "alter table uploaded_documents add column extracted_grades text not null default '[]'",
      "extraction_error": "alter table uploaded_documents add column extraction_error text",
      "extracted_at": "alter table uploaded_documents add column extracted_at text",
    }
    for column, statement in document_migrations.items():
      if column not in existing_columns:
        connection.execute(statement)
    connection.commit()


init_database()


def check_sqlite_ready() -> bool:
  try:
    with get_db_connection() as connection:
      connection.execute("select 1").fetchone()
    return True
  except Exception:
    return False


def check_upload_storage_ready() -> bool:
  try:
    UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)
    probe = UPLOAD_ROOT / f".health-{uuid.uuid4().hex}"
    probe.write_text("ok", encoding="utf-8")
    probe.unlink(missing_ok=True)
    return True
  except Exception:
    return False


def check_data_backend_ready() -> bool:
  if using_supabase():
    try:
      supabase_request(
        "GET",
        supabase_table_path("runtime_app_state", "select=state_key&limit=1"),
      )
      return True
    except Exception:
      return False
  return check_sqlite_ready()


def check_document_storage_ready() -> bool:
  if using_supabase():
    try:
      supabase_request(
        "GET",
        f"/storage/v1/bucket/{quote(get_supabase_storage_bucket(), safe='')}",
        content_type=None,
      )
      return True
    except Exception:
      return False
  return check_upload_storage_ready()


def sanitize_storage_segment(value: str, fallback: str = "item") -> str:
  cleaned = re.sub(r"[^a-zA-Z0-9_.-]", "-", value or "").strip(".-")
  return cleaned[:80] or fallback


def sanitize_upload_filename(filename: str) -> str:
  name = Path(filename or "document").name
  cleaned = re.sub(r"[^a-zA-Z0-9_. -]", "_", name).strip(" .")
  return cleaned[:120] or "document"


def validate_file_signature(content: bytes, extension: str) -> bool:
  """Validate file type via magic number signatures (first bytes of file)."""
  if len(content) < 4:
    return False
  
  magic_signatures = {
    ".pdf": b"%PDF",
    ".docx": b"PK\x03\x04",  # ZIP signature
    ".xlsx": b"PK\x03\x04",  # ZIP signature
    ".jpg": (b"\xFF\xD8\xFF",),
    ".jpeg": (b"\xFF\xD8\xFF",),
    ".png": b"\x89PNG\r\n\x1a\n",
  }
  
  if extension not in magic_signatures:
    return True  # Unknown extension - let MIME check handle it
  
  expected = magic_signatures[extension]
  
  if isinstance(expected, tuple):
    # For formats with multiple possible signatures
    return any(content.startswith(sig) for sig in expected)
  else:
    return content.startswith(expected)


def validate_mime_type(content_type: str | None, extension: str) -> bool:
  """Validate MIME type against file extension."""
  if not content_type:
    return False
  
  # Normalize MIME type
  mime_base = content_type.split(";")[0].lower().strip()
  
  allowed_mimes = {
    ".pdf": {"application/pdf"},
    ".docx": {
      "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
      "application/msword",
      "application/vnd.ms-word",
    },
    ".xlsx": {
      "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
      "application/vnd.ms-excel",
      "application/x-msexcel",
    },
    ".jpg": {"image/jpeg"},
    ".jpeg": {"image/jpeg"},
    ".png": {"image/png"},
  }
  
  if extension not in allowed_mimes:
    return True  # Unknown extension - accept any MIME
  
  return mime_base in allowed_mimes[extension]


def row_get(row: sqlite3.Row | dict[str, Any], key: str, default: Any = None) -> Any:
  if isinstance(row, sqlite3.Row):
    return row[key] if key in row.keys() else default
  if isinstance(row, dict):
    return row.get(key, default)
  return default


def parse_jsonish(value: Any, fallback: Any = None) -> Any:
  if value is None:
    return fallback
  if isinstance(value, (dict, list)):
    return value
  if isinstance(value, str):
    try:
      return json.loads(value)
    except json.JSONDecodeError:
      return fallback
  return fallback


def sanitize_guidance_text(value: Any, limit: int) -> str:
  text = str(value or "").replace("\x00", "").strip()
  return text[:limit]


def sanitize_profile_payload(payload: AuthProfileRequest) -> AuthProfileRequest:
  safe_name = sanitize_guidance_text(payload.name, 120)
  safe_district = sanitize_guidance_text(payload.district, 120)
  safe_stream = sanitize_guidance_text(payload.stream, 80)
  safe_leaving_year = sanitize_guidance_text(payload.leavingYear, 20)
  safe_income_band = payload.incomeBand if payload.incomeBand in {"low", "mid", "high"} else "mid"
  safe_need_signals = [sanitize_guidance_text(item, 60) for item in (payload.needSignals or [])[:10] if str(item or "").strip()]
  safe_preference_text = sanitize_guidance_text(payload.preferenceText, 2000)
  safe_grades = {
    sanitize_guidance_text(key, 40): sanitize_guidance_text(value, 10)
    for key, value in list((payload.grades or {}).items())[:50]
    if str(key or "").strip() and str(value or "").strip()
  }
  safe_documents = [sanitize_guidance_object(item, 0, 1800) for item in (payload.documents or [])[:25]]
  safe_shortlist = [sanitize_guidance_text(item, 120) for item in (payload.shortlist or [])[:25] if str(item or "").strip()]
  safe_pathways = {
    sanitize_guidance_text(key, 120): value
    for key, value in list((payload.shortlistPathways or {}).items())[:25]
    if str(key or "").strip() and value in {"primary", "backup", "considering", "not_interested"}
  }
  return AuthProfileRequest(
    name=safe_name,
    district=safe_district,
    stream=safe_stream,
    leavingYear=safe_leaving_year,
    incomeBand=safe_income_band,
    needSignals=safe_need_signals,
    preferenceText=safe_preference_text,
    grades=safe_grades,
    documents=safe_documents,
    shortlist=safe_shortlist,
    shortlistPathways=safe_pathways,
  )


def sanitize_guidance_object(value: Any, depth: int = 0, limit: int = 2000) -> Any:
  if depth > 4:
    return None
  if isinstance(value, str):
    return sanitize_guidance_text(value, limit)
  if isinstance(value, bool):
    return value
  if isinstance(value, (int, float)):
    return value
  if value is None:
    return None
  if isinstance(value, list):
    return [sanitize_guidance_object(item, depth + 1) for item in value[:50]]
  if isinstance(value, dict):
    cleaned: dict[str, Any] = {}
    for key, item in list(value.items())[:50]:
      safe_key = re.sub(r"[^a-zA-Z0-9_\-]", "_", str(key)).strip("_")
      if not safe_key:
        continue
      cleaned[safe_key] = sanitize_guidance_object(item, depth + 1)
    return cleaned
  return sanitize_guidance_text(value, limit)


def validate_and_sanitize_guidance_payload(payload: GuidanceRequest) -> GuidanceRequest:
  serialized = json.dumps(payload.model_dump(mode="json"), ensure_ascii=False, default=str)
  if len(serialized.encode("utf-8")) > AI_GUIDANCE_MAX_PAYLOAD_BYTES:
    raise HTTPException(status_code=413, detail="AI payload exceeds the supported size limit.")

  question = sanitize_guidance_text(payload.question, AI_GUIDANCE_MAX_QUESTION_LENGTH)
  if len(question.encode("utf-8")) > AI_GUIDANCE_MAX_QUESTION_LENGTH:
    raise HTTPException(status_code=400, detail="AI question is too long.")

  matches = [sanitize_guidance_object(item) for item in (payload.matches or [])[:AI_GUIDANCE_MAX_MATCHES]]
  blocked_matches = [sanitize_guidance_object(item) for item in (payload.blockedMatches or [])[:AI_GUIDANCE_MAX_BLOCKED_MATCHES]]
  documents = [sanitize_guidance_object(item) for item in (payload.documents or [])[:AI_GUIDANCE_MAX_DOCUMENTS]]
  conversation = [
    {
      "id": sanitize_guidance_text(item.get("id", ""), 80),
      "role": sanitize_guidance_text(item.get("role", "user"), 20),
      "content": sanitize_guidance_text(item.get("content", ""), AI_GUIDANCE_MAX_MESSAGE_LENGTH),
    }
    for item in (payload.conversation or [])[:AI_GUIDANCE_MAX_CONVERSATION_ITEMS]
    if isinstance(item, dict)
  ]

  sanitized = GuidanceRequest(
    profile=sanitize_guidance_object(payload.profile, 0, 5000) if isinstance(payload.profile, dict) else {},
    readiness=sanitize_guidance_object(payload.readiness, 0, 5000) if isinstance(payload.readiness, dict) else {},
    matches=matches,
    blockedMatches=blocked_matches,
    documents=documents,
    conversation=conversation,
    question=question,
    mode=payload.mode if payload.mode in {"guidance", "compare", "interview"} else "guidance",
  )
  if not sanitized.matches and not sanitized.blockedMatches:
    return sanitized
  return sanitized


def document_response(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
  extracted_grades = parse_jsonish(row_get(row, "extracted_grades"), [])
  extraction_text = row_get(row, "extraction_text")
  document_id = row_get(row, "id")
  try:
    if not isinstance(extracted_grades, list):
      extracted_grades = []
  except TypeError:
    extracted_grades = []
  return {
    "id": document_id,
    "userId": row_get(row, "user_id"),
    "name": row_get(row, "original_name"),
    "storedName": row_get(row, "stored_name"),
    "contentType": row_get(row, "content_type"),
    "size": row_get(row, "size_bytes"),
    "status": row_get(row, "status"),
    "uploadedAt": row_get(row, "uploaded_at"),
    "url": f"/api/documents/{document_id}/download",
    "extractionStatus": row_get(row, "extraction_status", "pending"),
    "extractedGrades": extracted_grades,
    "extractedTextPreview": (extraction_text or "")[:500],
    "extractionError": row_get(row, "extraction_error"),
    "extractedAt": row_get(row, "extracted_at"),
  }


DEVELOPER_PROMPT = """
You are EduGuide LS, an academic guidance assistant for students in Lesotho.
Use only the provided matcher payload. Do not invent institutions, fees, requirements,
sponsorship decisions, or official admissions outcomes.
Respect the matcher tier: "qualified" means currently meets captured rules, "almost"
means close or missing a small requirement, and "explore" means interest fit only.
Never upgrade a programme beyond the tier supplied by the matcher. When match.scoreBreakdown is present, use it to explain where the final match percentage came from; do not recalculate or change the supplied score. When match.requirementEvidence is present, use it to explain met requirements, warnings, and hard-gate failures. Hard-gate failures override a high percentage and must be described as blockers. Only put
qualified or almost programmes in top_recommendations and alternative routes;
use explore programmes for context only. Do not recommend science, technology,
engineering, health-science, or architecture pathways unless
the matcher payload already includes them as realistic options. Missing Mathematics,
Physical Science, Biology, or another hard subject gate must be treated as a blocker,
not as something the AI can overlook because the student is interested.
Respect supplied fundingBreakdown.policy values. For National University of Lesotho
and IEMS records, treat diploma/certificate or other non-degree programmes as not
NMDS sponsorship-ready when the payload says "degree+ only"; do not describe those
non-degree NUL pathways as funded options. Degree-and-higher NUL pathways may be
discussed as possible funding routes, but keep the estimate cautious and competitive.
Uploaded documents may include OCR/text extraction metadata. Treat extracted grades
as machine-read suggestions until the student applies or confirms them in the grade
form. Do not present extracted grades as an official transcript interpretation.
When application/source links and application document checklists are supplied in
the payload, use them for practical next steps. Do not invent deadlines; say
deadline tracking is not active yet unless the payload provides a verified date.

Answer the student's question when one is provided. Use the recent conversation
only for context and continuity; the current matcher payload remains the source
of truth. If mode is "compare", compare the strongest realistic qualified/almost
options instead of repeating generic advice. If mode is "interview", explain the
profile built from the interview answers, state what is still missing, and guide
the student toward the strongest qualified/almost matches produced by the matcher.
If the student asks why they are
blocked, explain the exact blocker from blocked_matches or requirement_gaps and
give a realistic next step. Ask one or two follow-up questions whenever important
profile details are missing.

Return concise JSON with exactly these keys:
- summary: string
- direct_answer: string
- top_recommendations: array of objects with programme, institution, tier, evidence, why, caution, action
- comparison: array of objects with programme, institution, tier, evidence, strength, concern
- study_plan: array of short strings
- document_checklist: array of short strings
- scholarship_note: string
- next_questions: array of short strings

Be warm, direct, and practical. If data confidence is low or a requirement is marked
under review, say the student should verify it with the institution/admin.
Every recommendation must cite evidence from the supplied programme object,
such as requirements, reasons, cautions, requirement_gaps, source, application,
or applicationDocuments. If evidence is weak, say so instead of sounding certain.
"""


def get_secret(name: str, *placeholders: str) -> str | None:
  value = os.getenv(name, "").strip()
  lower_value = value.lower()
  if not value or value in placeholders or lower_value.startswith("replace_with") or lower_value.startswith("your_"):
    return None
  return value


def get_gemini_api_key() -> str | None:
  return get_secret("GEMINI_API_KEY", "replace_with_your_gemini_api_key")


def get_openai_api_key() -> str | None:
  return get_secret("OPENAI_API_KEY", "replace_with_your_openai_api_key")


def get_brevo_api_key() -> str | None:
  return get_secret("BREVO_API_KEY", "replace_with_your_brevo_api_key")


def get_bool_env(name: str, default: bool = False) -> bool:
  value = os.getenv(name)
  if value is None:
    return default
  return value.strip().lower() in {"1", "true", "yes", "on"}


def email_transport() -> str:
  if get_brevo_api_key():
    return "brevo_api"
  if get_secret("SMTP_HOST"):
    return "smtp"
  return "none"


def smtp_missing_keys() -> list[str]:
  if get_brevo_api_key():
    return [] if get_secret("SMTP_FROM_EMAIL") else ["SMTP_FROM_EMAIL"]

  missing = []
  smtp_host = get_secret("SMTP_HOST")
  username = get_secret("SMTP_USERNAME")
  password = get_secret("SMTP_PASSWORD")
  for key in ["SMTP_HOST", "SMTP_FROM_EMAIL"]:
    if not get_secret(key):
      missing.append(key)
  if smtp_host and ("brevo" in smtp_host.lower() or username or password):
    for key in ["SMTP_USERNAME", "SMTP_PASSWORD"]:
      if not get_secret(key):
        missing.append(key)
  return missing


def smtp_configured() -> bool:
  return not smtp_missing_keys()


def email_debug_codes_enabled() -> bool:
  return get_bool_env("EMAIL_DEBUG_CODES", False)


def send_brevo_api_email(email: str, subject: str, body_lines: list[str]) -> None:
  api_key = get_brevo_api_key()
  from_email = get_secret("SMTP_FROM_EMAIL")
  if not api_key or not from_email:
    raise RuntimeError("BREVO_API_KEY and SMTP_FROM_EMAIL are required for Brevo API email delivery.")

  from_name = os.getenv("SMTP_FROM_NAME", "EduGuide LS").strip() or "EduGuide LS"
  payload = {
    "sender": {"name": from_name, "email": from_email},
    "to": [{"email": email}],
    "subject": subject,
    "textContent": "\n".join(body_lines),
  }
  request = urlrequest.Request(
    "https://api.brevo.com/v3/smtp/email",
    data=json.dumps(payload).encode("utf-8"),
    headers={
      "accept": "application/json",
      "api-key": api_key,
      "content-type": "application/json",
    },
    method="POST",
  )
  try:
    with urlrequest.urlopen(request, timeout=20) as response:
      response.read()
  except urlerror.HTTPError as exc:
    details = exc.read().decode("utf-8", errors="replace")[:600]
    raise RuntimeError(f"Brevo API returned {exc.code}: {details}") from exc


def send_plain_email(email: str, subject: str, body_lines: list[str]) -> None:
  if get_brevo_api_key():
    send_brevo_api_email(email, subject, body_lines)
    return

  smtp_host = get_secret("SMTP_HOST")
  from_email = get_secret("SMTP_FROM_EMAIL")
  missing = smtp_missing_keys()
  if missing:
    raise RuntimeError(f"Missing SMTP setting(s): {', '.join(missing)}.")

  smtp_port = int(os.getenv("SMTP_PORT", "465" if get_bool_env("SMTP_USE_SSL", False) else "587"))
  from_name = os.getenv("SMTP_FROM_NAME", "EduGuide LS").strip() or "EduGuide LS"
  use_ssl = get_bool_env("SMTP_USE_SSL", False)
  use_tls = get_bool_env("SMTP_USE_TLS", not use_ssl)
  username = get_secret("SMTP_USERNAME")
  password = get_secret("SMTP_PASSWORD")

  message = EmailMessage()
  message["Subject"] = subject
  message["From"] = f"{from_name} <{from_email}>"
  message["To"] = email
  message.set_content("\n".join(body_lines))

  if use_ssl:
    with smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=20) as server:
      if username and password:
        server.login(username, password)
      server.send_message(message)
    return

  with smtplib.SMTP(smtp_host, smtp_port, timeout=20) as server:
    if use_tls:
      server.starttls()
    if username and password:
      server.login(username, password)
    server.send_message(message)


def send_verification_email(email: str, name: str, code: str) -> None:
  send_plain_email(
    email,
    "Your EduGuide LS verification code",
    [
      f"Hello {name or 'student'},",
      "",
      f"Your EduGuide LS verification code is: {code}",
      f"It expires in {EMAIL_VERIFICATION_TTL_MINUTES} minutes.",
      "",
      "If you did not request this account, you can ignore this email.",
      "",
      "EduGuide LS",
    ],
  )


def send_password_reset_email(email: str, name: str, code: str) -> None:
  send_plain_email(
    email,
    "Reset your EduGuide LS password",
    [
      f"Hello {name or 'student'},",
      "",
      f"Your EduGuide LS password reset code is: {code}",
      f"This code expires in {EMAIL_VERIFICATION_TTL_MINUTES} minutes.",
      "",
      "If you did not request this, you can ignore this email.",
      "",
      "EduGuide LS",
    ],
  )


def get_supabase_url() -> str | None:
  value = get_secret("SUPABASE_URL", "replace_with_your_supabase_url")
  return value.rstrip("/") if value else None


def get_supabase_service_key() -> str | None:
  return get_secret(
    "SUPABASE_SERVICE_ROLE_KEY",
    "replace_with_your_supabase_service_role_key",
  ) or get_secret("SUPABASE_SERVICE_KEY", "replace_with_your_supabase_service_role_key")


def supabase_configured() -> bool:
  return bool(get_supabase_url() and get_supabase_service_key())


def get_data_backend() -> str:
  requested = os.getenv("DATA_BACKEND", "auto").strip().lower()
  if requested == "sqlite":
    return "sqlite"
  if requested == "supabase" and not supabase_configured():
    return "sqlite"
  return "supabase" if supabase_configured() else "sqlite"


def using_supabase() -> bool:
  return get_data_backend() == "supabase"


def get_supabase_storage_bucket() -> str:
  return (os.getenv("SUPABASE_STORAGE_BUCKET") or DEFAULT_SUPABASE_STORAGE_BUCKET).strip() or DEFAULT_SUPABASE_STORAGE_BUCKET


def supabase_headers(content_type: str | None = "application/json", prefer: str | None = None) -> dict[str, str]:
  key = get_supabase_service_key()
  if not key:
    raise RuntimeError("SUPABASE_SERVICE_ROLE_KEY is not configured.")
  headers = {
    "apikey": key,
    "Authorization": f"Bearer {key}",
    "Accept": "application/json",
  }
  if content_type:
    headers["Content-Type"] = content_type
  if prefer:
    headers["Prefer"] = prefer
  return headers


def supabase_request(
  method: str,
  path: str,
  body: Any = None,
  *,
  content_type: str | None = "application/json",
  prefer: str | None = None,
  extra_headers: dict[str, str] | None = None,
  raw: bool = False,
  timeout: int = 30,
) -> Any:
  base_url = get_supabase_url()
  if not base_url:
    raise RuntimeError("SUPABASE_URL is not configured.")
  data = None
  if body is not None:
    data = body if isinstance(body, bytes) else json.dumps(body, ensure_ascii=False).encode("utf-8")
  headers = supabase_headers(content_type, prefer)
  if extra_headers:
    headers.update(extra_headers)
  request = urlrequest.Request(
    f"{base_url}{path}",
    data=data,
    method=method,
    headers=headers,
  )
  try:
    with urlrequest.urlopen(request, timeout=timeout) as response:
      payload = response.read()
      if raw:
        return payload, dict(response.headers)
      if not payload:
        return None
      return json.loads(payload.decode("utf-8"))
  except urlerror.HTTPError as exc:
    details = exc.read().decode("utf-8", errors="ignore")
    raise RuntimeError(f"Supabase {method} {path} failed with {exc.code}: {details}") from exc


def supabase_table_path(table: str, query: str = "") -> str:
  suffix = f"?{query}" if query else ""
  return f"/rest/v1/{table}{suffix}"


def supabase_filter_value(value: str) -> str:
  return quote(value, safe="")


def get_ai_provider() -> str:
  provider = os.getenv("AI_PROVIDER", "gemini").strip().lower()
  return provider if provider in {"gemini", "openai"} else "gemini"


def guess_mime_type(path: Path, fallback: str | None = None) -> str:
  return fallback or mimetypes.guess_type(path.name)[0] or "application/octet-stream"


def cached_file_response(path: Path, media_type: str | None = None, headers: dict[str, str] | None = None, filename: str | None = None) -> FileResponse:
  return FileResponse(path, media_type=media_type or guess_mime_type(path), headers=headers or STATIC_CACHE_HEADERS, filename=filename)


def normalize_ocr_text(text: str) -> str:
  return re.sub(r"[ \t]+", " ", (text or "").replace("\r", "\n")).strip()


def extract_text_locally(path: Path, content_type: str | None = None) -> str:
  suffix = path.suffix.lower()
  mime_type = guess_mime_type(path, content_type)
  if suffix == ".pdf" or mime_type == "application/pdf":
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return normalize_ocr_text("\n".join(page.extract_text() or "" for page in reader.pages))
  if suffix == ".docx" or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
    from docx import Document

    document = Document(str(path))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
      for row in table.rows:
        lines.append(" | ".join(cell.text for cell in row.cells))
    return normalize_ocr_text("\n".join(lines))
  if suffix in {".txt", ".csv", ".tsv"} or mime_type.startswith("text/"):
    return normalize_ocr_text(path.read_text(encoding="utf-8", errors="ignore"))
  return ""


def extract_text_with_gemini_ocr(path: Path, content_type: str | None = None) -> str:
  api_key = get_gemini_api_key()
  if not api_key:
    raise RuntimeError("Gemini API key is not configured for OCR.")

  from google import genai
  from google.genai import types

  mime_type = guess_mime_type(path, content_type)
  file_part = types.Part.from_bytes(data=path.read_bytes(), mime_type=mime_type)
  client = genai.Client(api_key=api_key)
  last_error: Exception | None = None
  for model in get_gemini_model_candidates():
    try:
      response = client.models.generate_content(
        model=model,
        contents=[
          "Extract all visible text from this results slip, transcript, or academic document. Return plain text only. Preserve subject names and grade symbols such as A*, A, B, C, D, E, F, G, X, and Z.",
          file_part,
        ],
      )
      return normalize_ocr_text(response.text or "")
    except Exception as exc:
      last_error = exc
      if "NOT_FOUND" not in str(exc) and "not found" not in str(exc).lower() and "no longer available" not in str(exc).lower():
        raise
  raise RuntimeError(f"Gemini OCR failed for available model fallbacks: {last_error}")


def should_use_vision_ocr(path: Path, content_type: str | None, local_text: str) -> bool:
  suffix = path.suffix.lower()
  mime_type = guess_mime_type(path, content_type)
  if mime_type.startswith("image/"):
    return True
  if suffix == ".pdf" or mime_type == "application/pdf":
    return len(local_text) < 80
  return False


def grade_pattern() -> re.Pattern[str]:
  return re.compile(r"(?<![A-Z0-9])(?:A\*|[A-GXZ])(?![A-Z0-9*])", re.IGNORECASE)


def grade_from_nearby_text(text: str, alias: str) -> tuple[str | None, int | None]:
  lower = text.lower()
  alias_index = lower.find(alias.lower())
  if alias_index < 0:
    return None, None

  candidates = []
  for match in grade_pattern().finditer(text.upper()):
    distance = min(abs(match.start() - alias_index), abs(match.end() - (alias_index + len(alias))))
    if distance <= 90:
      candidates.append((distance, match.group(0).upper()))
  if not candidates:
    return None, None
  candidates.sort(key=lambda item: item[0])
  return candidates[0][1], candidates[0][0]


def parse_extracted_grades(text: str) -> list[dict[str, Any]]:
  extracted: dict[str, dict[str, Any]] = {}
  lines = [line.strip(" |:-") for line in normalize_ocr_text(text).split("\n") if line.strip()]
  searchable_units = lines + [normalize_ocr_text(text)]

  for subject in SUBJECT_ALIASES:
    for alias in sorted(subject["aliases"], key=len, reverse=True):
      for unit in searchable_units:
        if alias.lower() not in unit.lower():
          continue
        grade, distance = grade_from_nearby_text(unit, alias)
        if grade not in GRADE_VALUES:
          continue
        confidence = 0.88 if unit in lines and (distance or 99) <= 35 else 0.68
        current = extracted.get(subject["code"])
        if current and current["confidence"] >= confidence:
          continue
        extracted[subject["code"]] = {
          "code": subject["code"],
          "subject": subject["subject"],
          "grade": grade,
          "confidence": confidence,
          "evidence": unit[:180],
        }
  return sorted(extracted.values(), key=lambda item: item["subject"])


def encoded_storage_path(storage_path: str) -> str:
  return "/".join(quote(part, safe="") for part in storage_path.split("/") if part)


def upload_supabase_storage_object(storage_path: str, content: bytes, content_type: str | None) -> None:
  supabase_request(
    "POST",
    f"/storage/v1/object/{get_supabase_storage_bucket()}/{encoded_storage_path(storage_path)}",
    content,
    content_type=content_type or "application/octet-stream",
    extra_headers={"x-upsert": "true"},
    raw=True,
    timeout=60,
  )


def download_supabase_storage_object(storage_path: str) -> bytes:
  content, _headers = supabase_request(
    "GET",
    f"/storage/v1/object/{get_supabase_storage_bucket()}/{encoded_storage_path(storage_path)}",
    content_type=None,
    raw=True,
    timeout=60,
  )
  return content


def delete_supabase_storage_object(storage_path: str) -> None:
  try:
    supabase_request(
      "DELETE",
      f"/storage/v1/object/{get_supabase_storage_bucket()}/{encoded_storage_path(storage_path)}",
      content_type=None,
      raw=True,
      timeout=30,
    )
  except Exception:
    pass


def insert_document_record(record: dict[str, Any]) -> sqlite3.Row | dict[str, Any]:
  if using_supabase():
    rows = supabase_request(
      "POST",
      supabase_table_path("runtime_uploaded_documents"),
      record,
      prefer="return=representation",
    )
    return rows[0] if rows else record

  with get_db_connection() as connection:
    connection.execute(
      """
      insert into uploaded_documents (
        id, user_id, original_name, stored_name, content_type, size_bytes, storage_path, status
      )
      values (?, ?, ?, ?, ?, ?, ?, ?)
      """,
      (
        record["id"],
        record["user_id"],
        record["original_name"],
        record["stored_name"],
        record.get("content_type"),
        record["size_bytes"],
        record["storage_path"],
        record.get("status") or "Uploaded - OCR pending",
      ),
    )
    connection.commit()
    return connection.execute("select * from uploaded_documents where id = ?", (record["id"],)).fetchone()


def get_document_record(document_id: str) -> sqlite3.Row | dict[str, Any] | None:
  if using_supabase():
    rows = supabase_request(
      "GET",
      supabase_table_path(
        "runtime_uploaded_documents",
        f"id=eq.{supabase_filter_value(document_id)}&select=*&limit=1",
      ),
    )
    return rows[0] if rows else None

  with get_db_connection() as connection:
    return connection.execute("select * from uploaded_documents where id = ?", (document_id,)).fetchone()


def list_document_records(user_id: str) -> list[sqlite3.Row | dict[str, Any]]:
  if using_supabase():
    rows = supabase_request(
      "GET",
      supabase_table_path(
        "runtime_uploaded_documents",
        f"user_id=eq.{supabase_filter_value(user_id)}&select=*&order=uploaded_at.desc",
      ),
    )
    return rows or []

  with get_db_connection() as connection:
    return connection.execute(
      """
      select *
      from uploaded_documents
      where user_id = ?
      order by uploaded_at desc
      """,
      (user_id,),
    ).fetchall()


def delete_document_record(document_id: str) -> sqlite3.Row | dict[str, Any] | None:
  if using_supabase():
    rows = supabase_request(
      "DELETE",
      supabase_table_path(
        "runtime_uploaded_documents",
        f"id=eq.{supabase_filter_value(document_id)}&select=*",
      ),
      prefer="return=representation",
    )
    return rows[0] if rows else None

  with get_db_connection() as connection:
    row = connection.execute("select * from uploaded_documents where id = ?", (document_id,)).fetchone()
    if row:
      connection.execute("delete from uploaded_documents where id = ?", (document_id,))
      connection.commit()
    return row


def update_document_extraction(document_id: str, status: str, text: str = "", grades: list[dict[str, Any]] | None = None, error: str | None = None) -> sqlite3.Row | dict[str, Any]:
  final_status = "Uploaded - OCR extracted" if grades else "Uploaded - OCR checked"
  if status == "failed":
    final_status = "Uploaded - OCR failed"
  elif status == "no_text":
    final_status = "Uploaded - no readable text detected"
  elif status == "no_grades":
    final_status = "Uploaded - no grades detected"

  if using_supabase():
    rows = supabase_request(
      "PATCH",
      supabase_table_path(
        "runtime_uploaded_documents",
        f"id=eq.{supabase_filter_value(document_id)}&select=*",
      ),
      {
        "status": final_status,
        "extraction_status": status,
        "extraction_text": text[:20000],
        "extracted_grades": grades or [],
        "extraction_error": error,
        "extracted_at": now_iso(),
      },
      prefer="return=representation",
    )
    if not rows:
      raise HTTPException(status_code=404, detail="Document not found.")
    return rows[0]

  with get_db_connection() as connection:
    connection.execute(
      """
      update uploaded_documents
      set status = ?,
        extraction_status = ?,
        extraction_text = ?,
        extracted_grades = ?,
        extraction_error = ?,
        extracted_at = current_timestamp
      where id = ?
      """,
      (final_status, status, text[:20000], json.dumps(grades or [], ensure_ascii=False), error, document_id),
    )
    connection.commit()
    return connection.execute("select * from uploaded_documents where id = ?", (document_id,)).fetchone()


def extract_document(document_id: str) -> sqlite3.Row | dict[str, Any]:
  row = get_document_record(document_id)
  if not row:
    raise HTTPException(status_code=404, detail="Document not found.")

  temporary_path: Path | None = None
  if using_supabase():
    try:
      suffix = Path(str(row_get(row, "original_name") or row_get(row, "stored_name") or "document")).suffix
      with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
        temp_file.write(download_supabase_storage_object(str(row_get(row, "storage_path"))))
        temporary_path = Path(temp_file.name)
      path = temporary_path
    except Exception as exc:
      return update_document_extraction(document_id, "failed", error=str(exc))
  else:
    path = ROOT / str(row_get(row, "storage_path"))

  if not path.exists():
    return update_document_extraction(document_id, "failed", error="Stored document file is missing.")

  try:
    local_error = None
    try:
      local_text = extract_text_locally(path, row_get(row, "content_type"))
    except Exception as exc:
      local_text = ""
      local_error = str(exc)
    if should_use_vision_ocr(path, row_get(row, "content_type"), local_text):
      text = extract_text_with_gemini_ocr(path, row_get(row, "content_type"))
    else:
      if local_error:
        raise RuntimeError(local_error)
      text = local_text
    if not text.strip():
      return update_document_extraction(document_id, "no_text")
    grades = parse_extracted_grades(text)
    return update_document_extraction(document_id, "extracted" if grades else "no_grades", text=text, grades=grades)
  except Exception as exc:
    return update_document_extraction(document_id, "failed", error=str(exc))
  finally:
    if temporary_path:
      temporary_path.unlink(missing_ok=True)


def compact_match(match: dict[str, Any]) -> dict[str, Any]:
  scores = match.get("scores") or match.get("match") or {}
  requirements = (match.get("requirements") or [])[:4]
  reasons = (scores.get("reasons") or match.get("reasons") or [])[:3]
  cautions = (scores.get("cautions") or match.get("cautions") or [])[:3]
  requirement_gaps = (scores.get("requirementGaps") or match.get("requirementGaps") or [])[:4]
  application = match.get("application") if isinstance(match.get("application"), dict) else {}
  application_documents = match.get("applicationDocuments") or match.get("application_documents") or []
  source = match.get("source")
  evidence = []
  if requirements:
    evidence.append(f"Captured requirements: {'; '.join(str(item) for item in requirements[:2])}")
  if reasons:
    evidence.append(f"Matcher reason: {reasons[0]}")
  if requirement_gaps:
    evidence.append(f"Requirement gap: {requirement_gaps[0]}")
  if cautions:
    evidence.append(f"Caution: {cautions[0]}")
  if source:
    evidence.append(f"Source: {source}")
  if application.get("link"):
    link = application.get("link") or {}
    evidence.append(f"Application/source link: {link.get('label') or link.get('url')}")
  if application_documents:
    labels = [str(item.get("label") if isinstance(item, dict) else item) for item in application_documents[:3]]
    evidence.append(f"Application documents: {', '.join(labels)}")
  return {
    "programme": match.get("title"),
    "institution": match.get("institution"),
    "duration": match.get("duration"),
    "faculty": match.get("faculty"),
    "level": match.get("level"),
    "source": source,
    "source_type": match.get("sourceType") or match.get("source_type"),
    "careers": (match.get("careers") or [])[:4],
    "skills": (match.get("skills") or [])[:4],
    "requirements": requirements,
    "application": application,
    "application_documents": application_documents[:6],
    "hard_gate_passed": scores.get("hardGatePassed") if scores.get("hardGatePassed") is not None else match.get("hard_gate_passed"),
    "hard_gate_failures": (scores.get("hardGateFailures") or match.get("hardGateFailures") or [])[:4],
    "evidence": evidence[:5],
    "scores": {
      "overall": scores.get("overall"),
      "academic": scores.get("academic"),
      "interest": scores.get("interest") or scores.get("interests"),
      "eligibility": scores.get("eligibility"),
      "funding": scores.get("funding"),
      "funding_breakdown": scores.get("fundingBreakdown") or scores.get("funding_breakdown"),
      "confidence": scores.get("confidence"),
      "priority": scores.get("priority"),
    },
    "tier": scores.get("tier") or match.get("tier"),
    "tier_label": scores.get("tierLabel") or match.get("tierLabel"),
    "requirement_gaps": requirement_gaps,
    "reasons": reasons,
    "cautions": cautions,
  }


def is_recommendable_match(match: dict[str, Any]) -> bool:
  tier = str(match.get("tier_label") or match.get("tier") or "").lower()
  return "qualified" in tier or "almost" in tier


def guidance_match_key(programme: Any, institution: Any) -> str:
  return re.sub(r"[^a-z0-9]+", " ", f"{programme or ''} {institution or ''}".lower()).strip()


def evidence_text(match: dict[str, Any]) -> str:
  evidence = [str(item) for item in match.get("evidence") or [] if str(item).strip()]
  if evidence:
    return " | ".join(evidence[:3])
  fallback = []
  if match.get("requirements"):
    fallback.append(f"Captured requirements: {'; '.join(str(item) for item in match['requirements'][:2])}")
  if match.get("source"):
    fallback.append(f"Source: {match['source']}")
  if match.get("requirement_gaps"):
    fallback.append(f"Requirement gap: {match['requirement_gaps'][0]}")
  return " | ".join(fallback[:3]) or "Evidence is limited; verify this programme with the institution before applying."


def recommendation_from_match(match: dict[str, Any]) -> dict[str, Any]:
  caution_parts = match.get("cautions") or match.get("requirement_gaps") or []
  return {
    "programme": match.get("programme"),
    "institution": match.get("institution"),
    "tier": match.get("tier_label") or match.get("tier") or "Match",
    "evidence": evidence_text(match),
    "why": "; ".join(match.get("reasons") or ["This is supported by the matcher using your current grades, interests, and captured requirements."]),
    "caution": "; ".join(caution_parts[:2]) if caution_parts else "Verify final requirements with the institution before applying.",
    "action": "Open the programme or institution profile, check requirements and documents, then apply only through the captured official/source link.",
  }


def normalize_ai_guidance(guidance: dict[str, Any] | None, payload: GuidanceRequest) -> dict[str, Any]:
  response = guidance if isinstance(guidance, dict) else build_fallback(payload, "The AI response was empty.")
  compact_matches = [compact_match(item) for item in payload.matches[:8]]
  recommendable = [item for item in compact_matches if is_recommendable_match(item)]
  allowed = {guidance_match_key(item.get("programme"), item.get("institution")): item for item in recommendable}
  blocked_matches = [compact_match(item) for item in payload.blockedMatches[:6]]

  safe_recommendations = []
  for item in response.get("top_recommendations") or []:
    if not isinstance(item, dict):
      continue
    matched = allowed.get(guidance_match_key(item.get("programme"), item.get("institution")))
    if not matched:
      continue
    merged = recommendation_from_match(matched)
    merged.update({
      "why": str(item.get("why") or merged["why"])[:900],
      "caution": str(item.get("caution") or merged["caution"])[:900],
      "action": str(item.get("action") or merged["action"])[:900],
      "evidence": evidence_text(matched),
      "tier": matched.get("tier_label") or matched.get("tier") or merged["tier"],
    })
    safe_recommendations.append(merged)

  if not safe_recommendations and recommendable:
    safe_recommendations = [recommendation_from_match(item) for item in recommendable[:3]]

  safe_comparison = []
  for item in response.get("comparison") or []:
    if not isinstance(item, dict):
      continue
    matched = allowed.get(guidance_match_key(item.get("programme"), item.get("institution")))
    if not matched:
      continue
    safe_comparison.append({
      "programme": matched.get("programme"),
      "institution": matched.get("institution"),
      "tier": matched.get("tier_label") or matched.get("tier") or "Match",
      "evidence": evidence_text(matched),
      "strength": str(item.get("strength") or "; ".join(matched.get("reasons") or ["Supported by current matcher evidence."]))[:900],
      "concern": str(item.get("concern") or "; ".join(matched.get("cautions") or matched.get("requirement_gaps") or ["Confirm final entry requirements."]))[:900],
    })
  if not safe_comparison and safe_recommendations:
    safe_comparison = [
      {
        "programme": item["programme"],
        "institution": item["institution"],
        "tier": item["tier"],
        "evidence": item["evidence"],
        "strength": item["why"],
        "concern": item["caution"],
      }
      for item in safe_recommendations[:3]
    ]

  if not recommendable and blocked_matches:
    blocker = blocked_matches[0]
    gaps = blocker.get("requirement_gaps") or blocker.get("cautions") or blocker.get("hard_gate_failures") or []
    response["summary"] = "No qualified or almost-qualified recommendation is strong enough yet from the current captured evidence."
    response["direct_answer"] = (
      f"The closest blocked pathway is {blocker.get('programme')} at {blocker.get('institution')}. "
      f"Main blocker: {'; '.join(str(item) for item in gaps[:2]) or 'captured requirements are not met yet'}."
    )

  response["top_recommendations"] = safe_recommendations[:4]
  response["comparison"] = safe_comparison[:4]
  response["study_plan"] = [str(item)[:240] for item in (response.get("study_plan") or [])[:5]]
  response["document_checklist"] = [str(item)[:240] for item in (response.get("document_checklist") or [])[:8]]
  response["next_questions"] = [str(item)[:180] for item in (response.get("next_questions") or [])[:3]]
  response["summary"] = str(response.get("summary") or "Guidance generated from your current matcher evidence.")[:1200]
  response["direct_answer"] = str(response.get("direct_answer") or "Use qualified and almost-qualified matches first; blocked paths are preparation goals, not current application recommendations.")[:1600]
  response["scholarship_note"] = str(response.get("scholarship_note") or "Funding readiness is an estimate only; official sponsorship decisions remain with the sponsor/NMDS process.")[:1200]
  return response


def build_fallback(payload: GuidanceRequest, error: str | None = None) -> dict[str, Any]:
  compact_matches = [compact_match(item) for item in payload.matches[:8]]
  top_matches = [item for item in compact_matches if is_recommendable_match(item)][:3]
  explore_matches = [item for item in compact_matches if not is_recommendable_match(item)][:3]
  blocked_matches = [compact_match(item) for item in payload.blockedMatches[:5]]
  best = top_matches[0] if top_matches else (explore_matches[0] if explore_matches else {})
  programme = best.get("programme") or "your strongest current match"
  institution = best.get("institution") or "the matched institution"
  cautions = best.get("cautions") or []
  caution = cautions[0] if cautions else "Verify final requirements before applying."
  question = (payload.question or "").strip()
  blocked_summary = ""
  if blocked_matches:
    blocked = blocked_matches[0]
    gaps = blocked.get("requirement_gaps") or blocked.get("cautions") or []
    blocked_summary = f" A blocked example is {blocked.get('programme')} at {blocked.get('institution')}: {'; '.join(gaps[:2])}."
  uploaded_documents = [item for item in payload.documents if item.get("name")]
  extracted_grade_count = sum(len(item.get("extractedGrades") or []) for item in uploaded_documents)
  document_checklist = [
    "Certificates, results slip, or transcript",
    "National ID, passport, or birth certificate",
    "Admission, application, or registration evidence",
    "Bank account confirmation",
    "Residence/chief letter or guarantor evidence",
    "Prior NMDS, CHE evaluation, CV, or study leave if applicable",
  ]
  if uploaded_documents:
    uploaded_names = ", ".join(item["name"] for item in uploaded_documents[:3])
    more_note = " and more" if len(uploaded_documents) > 3 else ""
    document_checklist = [
      f"Uploaded: {uploaded_names}{more_note}",
      f"Machine-detected grade suggestions: {extracted_grade_count}. Confirm them in the grade form before relying on matches.",
      "Keep ID/birth certificate, bank confirmation, and admission/application evidence ready.",
      "If applicable, prepare residence/chief/guarantor evidence, prior NMDS documents, CHE evaluation, CV, or study leave letter.",
    ]

  response = {
    "summary": (
      f"Your strongest realistic match is {programme} at {institution}. The matcher should still be treated as guidance, not an admission decision."
      if top_matches
      else f"No qualified or almost-qualified recommendation is strong enough yet.{blocked_summary} Add/confirm grades and check requirement gaps before applying."
    ),
    "direct_answer": (
      f"For your question: {question} Based on the current evidence, use qualified/almost options first. {programme} at {institution} is the best available path to inspect, and blocked paths should only be used as future goals."
      if question
      else "Use the tier labels first: qualified options are strongest, almost options need small fixes or confirmation, and explore options are interest-fit pathways, not application recommendations."
    ),
    "top_recommendations": [
      {
        "programme": item.get("programme"),
        "institution": item.get("institution"),
        "tier": item.get("tier_label") or item.get("tier") or "Explore",
        "evidence": evidence_text(item),
        "why": "; ".join(item.get("reasons") or ["It aligns with your selected profile signals."]),
        "caution": "; ".join(item.get("cautions") or ["Confirm requirements with the institution."]),
        "action": "Compare requirements, duration, careers, and funding readiness before shortlisting.",
      }
      for item in top_matches
    ],
    "comparison": [
      {
        "programme": item.get("programme"),
        "institution": item.get("institution"),
        "tier": item.get("tier_label") or item.get("tier") or "Explore",
        "evidence": evidence_text(item),
        "strength": "; ".join(item.get("reasons") or ["Good profile alignment."]),
        "concern": "; ".join(item.get("requirement_gaps") or item.get("cautions") or ["Confirm final entry requirements."]),
      }
      for item in top_matches
    ],
    "study_plan": [
      "Keep Mathematics and English strong because many programmes use them for ranking.",
      "Focus revision on subjects connected to your top three matched pathways.",
      "Use blocked-pathway notes to identify missing or weak requirement areas.",
    ],
    "document_checklist": document_checklist,
    "scholarship_note": "The scholarship/NMDS-style score is only an estimate. Review academic fit, background/need, priority-programme alignment, and the sponsorship document checklist before treating a pathway as funding-ready. Final sponsorship decisions remain with the official portal and NMDS process.",
    "next_questions": [
      "Which of the qualified or almost-qualified matches feels most realistic for your marks?",
      "Do you want to compare careers, fees, or requirements first?",
      "Which missing document can you upload next?",
    ],
  }
  if error:
    response["summary"] = f"{response['summary']} AI service note: {error}"
  if caution and response["top_recommendations"]:
    response["top_recommendations"][0]["caution"] = caution
  return response


def parse_json_response(text: str) -> dict[str, Any] | None:
  cleaned = text.strip()
  fenced = re.search(r"```(?:json)?\s*(.*?)```", cleaned, flags=re.DOTALL | re.IGNORECASE)
  if fenced:
    cleaned = fenced.group(1).strip()
  try:
    return json.loads(cleaned)
  except json.JSONDecodeError:
    return None


def build_request_payload(payload: GuidanceRequest, conversation: list[dict[str, Any]] | None = None) -> dict[str, Any]:
  compact_matches = [compact_match(item) for item in payload.matches[:8]]
  recommendable_matches = [item for item in compact_matches if is_recommendable_match(item)]
  mode = payload.mode if payload.mode in {"guidance", "compare", "interview"} else "guidance"
  return {
    "task": {
      "mode": mode,
      "question": (payload.question or "").strip(),
    },
    "profile": payload.profile,
    "readiness": payload.readiness,
    "documents": payload.documents,
    "recommendable_matches": recommendable_matches[:6],
    "explore_matches": [item for item in compact_matches if not is_recommendable_match(item)][:4],
    "blocked_matches": [compact_match(item) for item in payload.blockedMatches[:6]],
    "conversation": [
      {
        "id": str(item.get("id", ""))[:80],
        "role": str(item.get("role", "user"))[:20],
        "content": str(item.get("content", ""))[:900],
      }
      for item in (conversation if conversation is not None else payload.conversation)[-AI_CHAT_CONTEXT_LIMIT:]
      if isinstance(item, dict) and str(item.get("content", "")).strip()
    ],
    "matches": compact_matches,
  }


def save_recommendation_run(payload: GuidanceRequest, request_payload: dict[str, Any]) -> None:
  try:
    if using_supabase():
      supabase_request(
        "POST",
        supabase_table_path("runtime_recommendation_runs"),
        {
          "mode": request_payload.get("task", {}).get("mode", "guidance"),
          "question": request_payload.get("task", {}).get("question") or None,
          "profile_name": (request_payload.get("profile") or {}).get("name"),
          "payload": request_payload,
        },
        prefer="return=minimal",
      )
      return

    with get_db_connection() as connection:
      connection.execute(
        """
        insert into recommendation_runs (mode, question, profile_name, payload)
        values (?, ?, ?, ?)
        """,
        (
          request_payload.get("task", {}).get("mode", "guidance"),
          request_payload.get("task", {}).get("question") or None,
          (request_payload.get("profile") or {}).get("name"),
          json.dumps(request_payload, ensure_ascii=False),
        ),
      )
      connection.commit()
  except Exception:
    pass


def sanitize_event_type(value: str | None) -> str:
  event_type = re.sub(r"[^a-zA-Z0-9_.:-]", "_", (value or "").strip().lower())
  return event_type[:80]


def sanitize_event_payload(value: Any, depth: int = 0) -> Any:
  if depth > 3:
    return None
  if isinstance(value, str):
    return value.strip()[:700]
  if isinstance(value, (int, float, bool)) or value is None:
    return value
  if isinstance(value, list):
    return [sanitize_event_payload(item, depth + 1) for item in value[:30]]
  if isinstance(value, dict):
    sanitized: dict[str, Any] = {}
    for key, item in list(value.items())[:50]:
      clean_key = re.sub(r"[^a-zA-Z0-9_.:-]", "_", str(key).strip())[:80]
      if clean_key:
        sanitized[clean_key] = sanitize_event_payload(item, depth + 1)
    return sanitized
  return str(value)[:300]


def insert_runtime_event(user_id: str, event_type: str, label: str | None = None, payload: dict[str, Any] | None = None) -> dict[str, Any]:
  safe_event_type = sanitize_event_type(event_type)
  if not safe_event_type:
    raise ValueError("Event type is required.")
  event = {
    "id": f"evt-{uuid.uuid4().hex[:18]}",
    "user_id": str(user_id),
    "event_type": safe_event_type,
    "label": (label or "").strip()[:240] or None,
    "payload": sanitize_event_payload(payload or {}),
    "created_at": now_iso(),
  }
  if using_supabase():
    supabase_request(
      "POST",
      supabase_table_path("runtime_events"),
      event,
      prefer="return=minimal",
    )
    return event

  with get_db_connection() as connection:
    connection.execute(
      """
      insert into runtime_events (id, user_id, event_type, label, payload, created_at)
      values (?, ?, ?, ?, ?, ?)
      """,
      (
        event["id"],
        event["user_id"],
        event["event_type"],
        event["label"],
        json.dumps(event["payload"], ensure_ascii=False),
        event["created_at"],
      ),
    )
    connection.commit()
  return event


def list_runtime_events(limit: int = 1000) -> list[dict[str, Any]]:
  capped_limit = max(1, min(limit, 5000))
  if using_supabase():
    return supabase_request(
      "GET",
      supabase_table_path(
        "runtime_events",
        f"select=id,user_id,event_type,label,payload,created_at&order=created_at.desc&limit={capped_limit}",
      ),
    ) or []

  with get_db_connection() as connection:
    rows = connection.execute(
      """
      select id, user_id, event_type, label, payload, created_at
      from runtime_events
      order by created_at desc
      limit ?
      """,
      (capped_limit,),
    ).fetchall()
  return [
    {
      **dict(row),
      "payload": parse_jsonish(row["payload"], {}),
    }
    for row in rows
  ]


def safe_list_runtime_events(limit: int = 1000) -> list[dict[str, Any]]:
  try:
    return list_runtime_events(limit)
  except Exception:
    return []


def safe_insert_runtime_event(user_id: str, event_type: str, label: str | None = None, payload: dict[str, Any] | None = None) -> dict[str, Any] | None:
  try:
    return insert_runtime_event(user_id, event_type, label, payload)
  except Exception:
    return None


def normalize_email(value: str | None) -> str:
  return (value or "").strip().lower()


def mask_email(value: str | None) -> str:
  email = normalize_email(value)
  if "@" not in email:
    return email
  name, domain = email.split("@", 1)
  visible = name[:2] if len(name) > 2 else name[:1]
  return f"{visible}{'*' * max(2, len(name) - len(visible))}@{domain}"


def hash_password(password: str, salt: str | None = None) -> tuple[str, str]:
  password_salt = salt or secrets.token_hex(16)
  digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), password_salt.encode("utf-8"), 180_000)
  return password_salt, digest.hex()


def verify_password(password: str, user: dict[str, Any]) -> bool:
  password_hash = user.get("passwordHash")
  password_salt = user.get("passwordSalt")
  if password_hash and password_salt:
    _, digest = hash_password(password, password_salt)
    return secrets.compare_digest(digest, str(password_hash))
  legacy_password = user.get("password")
  return bool(legacy_password) and secrets.compare_digest(str(legacy_password), password)


def load_state_payload(state_key: str, fallback: Any = None) -> Any:
  if using_supabase():
    rows = supabase_request(
      "GET",
      supabase_table_path(
        "runtime_app_state",
        f"state_key=eq.{supabase_filter_value(state_key)}&select=payload&limit=1",
      ),
    )
    if not rows:
      return fallback
    return parse_jsonish(rows[0].get("payload"), fallback)

  with get_db_connection() as connection:
    row = connection.execute("select payload from app_state where state_key = ?", (state_key,)).fetchone()
  if not row:
    return fallback
  try:
    return json.loads(row["payload"])
  except json.JSONDecodeError:
    return fallback


def save_state_payload(state_key: str, payload: Any) -> None:
  if using_supabase():
    supabase_request(
      "POST",
      supabase_table_path("runtime_app_state", "on_conflict=state_key"),
      {
        "state_key": state_key,
        "payload": payload,
        "updated_at": now_iso(),
      },
      prefer="resolution=merge-duplicates,return=minimal",
    )
    return

  serialized = json.dumps(payload, ensure_ascii=False)
  with get_db_connection() as connection:
    connection.execute(
      """
      insert into app_state (state_key, payload, updated_at)
      values (?, ?, current_timestamp)
      on conflict(state_key) do update set
        payload = excluded.payload,
        updated_at = current_timestamp
      """,
      (state_key, serialized),
    )
    connection.commit()


def list_state_payloads() -> list[dict[str, Any]]:
  if using_supabase():
    rows = supabase_request(
      "GET",
      supabase_table_path("runtime_app_state", "select=state_key,payload,updated_at&order=state_key.asc"),
    )
    return rows or []

  with get_db_connection() as connection:
    rows = connection.execute("select state_key, payload, updated_at from app_state").fetchall()
  return [
    {
      "state_key": row["state_key"],
      "payload": parse_jsonish(row["payload"], {}),
      "updated_at": row["updated_at"],
    }
    for row in rows
  ]


def delete_state_payload(state_key: str) -> int:
  if using_supabase():
    rows = supabase_request(
      "DELETE",
      supabase_table_path("runtime_app_state", f"state_key=eq.{supabase_filter_value(state_key)}&select=state_key"),
      prefer="return=representation",
    )
    return len(rows or [])

  with get_db_connection() as connection:
    cursor = connection.execute("delete from app_state where state_key = ?", (state_key,))
    connection.commit()
  return cursor.rowcount


def normalize_ai_chat_message(item: dict[str, Any] | sqlite3.Row | None) -> dict[str, Any] | None:
  if not item:
    return None
  role = "assistant" if row_get(item, "role") == "assistant" else "user"
  content = str(row_get(item, "content") or "").strip()
  if not content:
    return None
  message_id = sanitize_storage_segment(str(row_get(item, "id") or f"ai-{uuid.uuid4().hex[:12]}"), "ai-message")
  return {
    "id": message_id,
    "role": role,
    "content": content[:2400],
    "at": row_get(item, "at") or row_get(item, "created_at") or now_iso(),
  }


def merge_ai_chat_histories(*histories: list[dict[str, Any]]) -> list[dict[str, Any]]:
  merged: list[dict[str, Any]] = []
  seen: set[str] = set()
  for history in histories:
    for raw_message in history or []:
      message = normalize_ai_chat_message(raw_message)
      if not message:
        continue
      key = message["id"] if message.get("id") else f"{message['role']}:{message['content']}"
      fallback_key = f"{message['role']}:{message['content']}"
      if key in seen or fallback_key in seen:
        continue
      seen.add(key)
      seen.add(fallback_key)
      merged.append(message)
  return merged[-AI_CHAT_HISTORY_LIMIT:]


def list_ai_chat_history(user_id: str) -> list[dict[str, Any]]:
  if using_supabase():
    rows = supabase_request(
      "GET",
      supabase_table_path(
        "runtime_ai_chat_messages",
        f"user_id=eq.{supabase_filter_value(user_id)}&select=id,role,content,created_at&order=created_at.desc&limit={AI_CHAT_HISTORY_LIMIT}",
      ),
    ) or []
    return [message for message in (normalize_ai_chat_message(row) for row in reversed(rows)) if message]

  with get_db_connection() as connection:
    rows = connection.execute(
      """
      select id, role, content, created_at
      from ai_chat_messages
      where user_id = ?
      order by created_at desc
      limit ?
      """,
      (user_id, AI_CHAT_HISTORY_LIMIT),
    ).fetchall()
  return [message for message in (normalize_ai_chat_message(row) for row in reversed(rows)) if message]


def save_ai_chat_history(user_id: str, messages: list[dict[str, Any]]) -> list[dict[str, Any]]:
  normalized = merge_ai_chat_histories(messages)
  if using_supabase():
    supabase_request(
      "DELETE",
      supabase_table_path("runtime_ai_chat_messages", f"user_id=eq.{supabase_filter_value(user_id)}"),
      prefer="return=minimal",
    )
    if normalized:
      supabase_request(
        "POST",
        supabase_table_path("runtime_ai_chat_messages"),
        [
          {
            "id": message["id"],
            "user_id": user_id,
            "role": message["role"],
            "content": message["content"],
            "payload": {},
            "created_at": message.get("at") or now_iso(),
          }
          for message in normalized
        ],
        prefer="return=minimal",
      )
    return normalized

  with get_db_connection() as connection:
    connection.execute("delete from ai_chat_messages where user_id = ?", (user_id,))
    connection.executemany(
      """
      insert into ai_chat_messages (id, user_id, role, content, payload, created_at)
      values (?, ?, ?, ?, ?, ?)
      """,
      [
        (
          message["id"],
          user_id,
          message["role"],
          message["content"],
          "{}",
          message.get("at") or now_iso(),
        )
        for message in normalized
      ],
    )
    connection.commit()
  return normalized


def clear_ai_chat_history(user_id: str) -> None:
  if using_supabase():
    supabase_request(
      "DELETE",
      supabase_table_path("runtime_ai_chat_messages", f"user_id=eq.{supabase_filter_value(user_id)}"),
      prefer="return=minimal",
    )
    return

  with get_db_connection() as connection:
    connection.execute("delete from ai_chat_messages where user_id = ?", (user_id,))
    connection.commit()


def safe_list_ai_chat_history(user_id: str) -> list[dict[str, Any]]:
  try:
    return list_ai_chat_history(user_id)
  except Exception:
    return []


def safe_save_ai_chat_history(user_id: str, messages: list[dict[str, Any]]) -> list[dict[str, Any]]:
  try:
    return save_ai_chat_history(user_id, messages)
  except Exception:
    return merge_ai_chat_histories(messages)


def safe_clear_ai_chat_history(user_id: str) -> None:
  try:
    clear_ai_chat_history(user_id)
  except Exception:
    pass


def guidance_to_chat_text(guidance: dict[str, Any]) -> str:
  parts = [guidance.get("summary"), guidance.get("direct_answer"), guidance.get("scholarship_note")]
  recommendations = guidance.get("top_recommendations")
  if isinstance(recommendations, list) and recommendations:
    parts.append(
      "Top recommendations: "
      + " | ".join(
        f"{item.get('programme') or 'Programme'} at {item.get('institution') or 'Institution'}: {item.get('why') or item.get('action') or 'match'}"
        for item in recommendations[:3]
        if isinstance(item, dict)
      )
    )
  study_plan = guidance.get("study_plan")
  if isinstance(study_plan, list) and study_plan:
    parts.append("Study plan: " + " | ".join(str(item) for item in study_plan[:5]))
  next_questions = guidance.get("next_questions")
  if isinstance(next_questions, list) and next_questions:
    parts.append("Next questions: " + " | ".join(str(item) for item in next_questions[:3]))
  return "\n".join(str(part) for part in parts if part).strip()[:2400] or "Guidance generated from your current EduGuide profile."


def now_iso() -> str:
  return datetime.now(timezone.utc).isoformat(timespec="milliseconds").replace("+00:00", "Z")


def parse_iso_datetime(value: str | None) -> datetime | None:
  if not value:
    return None
  try:
    parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
  except ValueError:
    return None


def minutes_from_now(minutes: int) -> str:
  return (datetime.now(timezone.utc) + timedelta(minutes=minutes)).isoformat(timespec="milliseconds").replace("+00:00", "Z")


def get_request_ip(request: Request | None) -> str:
  if not request:
    return "unknown"
  forwarded_for = request.headers.get("x-forwarded-for", "")
  if forwarded_for:
    return forwarded_for.split(",", 1)[0].strip()[:80] or "unknown"
  return (request.client.host if request.client else "unknown")[:80]


def check_rate_limit(request: Request | None, bucket: str, limit: int, window_seconds: int, identifier: str | None = None) -> None:
  now = datetime.now(timezone.utc)
  cutoff = now - timedelta(seconds=window_seconds)
  key = f"{bucket}:{identifier or ''}:{get_request_ip(request)}"
  hits = [hit for hit in RATE_LIMIT_STATE.get(key, []) if hit > cutoff]
  if len(hits) >= limit:
    retry_after = max(1, int((hits[0] + timedelta(seconds=window_seconds) - now).total_seconds()))
    raise HTTPException(
      status_code=429,
      detail=f"Too many requests. Try again in {retry_after} seconds.",
      headers={"Retry-After": str(retry_after)},
    )
  hits.append(now)
  RATE_LIMIT_STATE[key] = hits
  if len(RATE_LIMIT_STATE) > 5000:
    stale_keys = [item_key for item_key, item_hits in RATE_LIMIT_STATE.items() if not any(hit > cutoff for hit in item_hits)]
    for stale_key in stale_keys[:1000]:
      RATE_LIMIT_STATE.pop(stale_key, None)


def cleanup_expired_security_records(force: bool = False) -> dict[str, int]:
  global LAST_SECURITY_CLEANUP_AT
  now = datetime.now(timezone.utc)
  if not force and LAST_SECURITY_CLEANUP_AT and now - LAST_SECURITY_CLEANUP_AT < timedelta(seconds=SECURITY_CLEANUP_INTERVAL_SECONDS):
    return {"sessions": 0, "codes": 0}
  LAST_SECURITY_CLEANUP_AT = now
  session_cutoff = (now - timedelta(days=AUTH_SESSION_TTL_DAYS)).isoformat(timespec="milliseconds").replace("+00:00", "Z")
  verification_cutoff = (now - timedelta(days=2)).isoformat(timespec="milliseconds").replace("+00:00", "Z")
  removed = {"sessions": 0, "codes": 0}
  try:
    if using_supabase():
      supabase_request(
        "DELETE",
        supabase_table_path("runtime_auth_sessions", f"last_seen_at=lt.{supabase_filter_value(session_cutoff)}"),
        prefer="return=minimal",
      )
      supabase_request(
        "DELETE",
        supabase_table_path("runtime_email_verifications", f"expires_at=lt.{supabase_filter_value(verification_cutoff)}"),
        prefer="return=minimal",
      )
      return removed

    with get_db_connection() as connection:
      cursor = connection.execute(
        "delete from auth_sessions where datetime(last_seen_at) < datetime(?)",
        (session_cutoff,),
      )
      removed["sessions"] = cursor.rowcount if cursor.rowcount and cursor.rowcount > 0 else 0
      cursor = connection.execute(
        "delete from email_verifications where datetime(expires_at) < datetime(?)",
        (verification_cutoff,),
      )
      removed["codes"] = cursor.rowcount if cursor.rowcount and cursor.rowcount > 0 else 0
      connection.commit()
  except Exception:
    return removed
  return removed


def maybe_cleanup_security_records() -> None:
  cleanup_expired_security_records(False)


def session_expired(last_seen_at: str | None, created_at: str | None = None) -> bool:
  reference = parse_iso_datetime(last_seen_at) or parse_iso_datetime(created_at)
  if not reference:
    return False
  return reference < datetime.now(timezone.utc) - timedelta(days=AUTH_SESSION_TTL_DAYS)


def add_user_activity(user: dict[str, Any], activity_type: str, label: str, actor: dict[str, Any] | None = None, metadata: dict[str, Any] | None = None) -> None:
  timestamp = now_iso()
  activity = user.setdefault("activity", [])
  activity.insert(
    0,
    {
      "id": f"act-{uuid.uuid4().hex[:12]}",
      "type": activity_type,
      "label": label,
      "at": timestamp,
      "actorId": (actor or user).get("id"),
      "actorName": (actor or user).get("name", "System"),
      "metadata": metadata or {},
    },
  )
  user["activity"] = activity[:45]
  user["lastActiveAt"] = timestamp
  user["lastActivity"] = label


def normalize_institution_name(value: Any) -> str:
  return re.sub(r"\s+", " ", str(value or "").strip())[:180]


def get_user_managed_institution(user: dict[str, Any] | None) -> str:
  if not user:
    return ""
  return normalize_institution_name(
    user.get("managedInstitution")
    or user.get("institution")
    or user.get("institutionName")
    or user.get("assignedInstitution")
  )


def institutions_match(left: Any, right: Any) -> bool:
  return normalize_institution_name(left).casefold() == normalize_institution_name(right).casefold()


def normalize_auth_user(user: dict[str, Any]) -> dict[str, Any] | None:
  email = normalize_email(user.get("email"))
  if not email or user.get("id") in {"demo-student", "demo-admin"}:
    return None
  created_at = user.get("createdAt") or now_iso()
  email_verified_at = user.get("emailVerifiedAt") or created_at
  role = user.get("role") if user.get("role") in {"owner", "admin", "student", "institution_admin"} else "student"
  managed_institution = get_user_managed_institution(user)
  return {
    "id": user.get("id") or f"user-{uuid.uuid4().hex[:12]}",
    "name": (user.get("name") or email).strip(),
    "email": email,
    "passwordHash": user.get("passwordHash"),
    "passwordSalt": user.get("passwordSalt"),
    "password": user.get("password"),
    "role": role,
    "managedInstitution": managed_institution if role == "institution_admin" else "",
    "status": user.get("status") or "active",
    "district": user.get("district") or "",
    "phone": user.get("phone") or "",
    "stream": user.get("stream") or "",
    "leavingYear": user.get("leavingYear") or "",
    "incomeBand": user.get("incomeBand") or "mid",
    "needSignals": user.get("needSignals") if isinstance(user.get("needSignals"), list) else [],
    "preferenceText": user.get("preferenceText") or "",
    "grades": user.get("grades") if isinstance(user.get("grades"), dict) else {},
    "documents": user.get("documents") if isinstance(user.get("documents"), list) else [],
    "shortlist": user.get("shortlist") if isinstance(user.get("shortlist"), list) else [],
    "shortlistPathways": user.get("shortlistPathways") if isinstance(user.get("shortlistPathways"), dict) else {},
    "createdAt": created_at,
    "emailVerifiedAt": email_verified_at,
    "reviewedAt": user.get("reviewedAt") or (created_at if role in {"owner", "admin", "institution_admin"} else None),
    "lastActiveAt": user.get("lastActiveAt") or created_at,
    "lastActivity": user.get("lastActivity") or "Account created",
    "activity": user.get("activity") if isinstance(user.get("activity"), list) else [],
  }


def get_auth_users_internal() -> list[dict[str, Any]]:
  payload = load_state_payload("auth_users", {"users": []})
  users = payload.get("users", []) if isinstance(payload, dict) else []
  normalized = []
  seen = set()
  for user in users:
    if not isinstance(user, dict):
      continue
    item = normalize_auth_user(user)
    if not item or item["email"] in seen:
      continue
    seen.add(item["email"])
    normalized.append(item)
  return normalized


def save_auth_users_internal(users: list[dict[str, Any]]) -> None:
  save_state_payload("auth_users", {"users": users})


def public_user(user: dict[str, Any]) -> dict[str, Any]:
  safe = {key: value for key, value in user.items() if key not in {"password", "passwordHash", "passwordSalt"}}
  return safe


INSTITUTION_PROPOSAL_STATE_KEY = "institution_proposals"
INSTITUTION_PROPOSAL_FIELDS = {
  "name",
  "code",
  "faculty",
  "category",
  "level",
  "duration",
  "deliveryMode",
  "requirementsSummary",
  "overview",
  "sourceUrl",
  "supportingSourcePath",
  "sourceNote",
  "feeNote",
  "supportingFeeSourcePath",
  "careers",
}
INSTITUTION_PROPOSAL_STATUSES = {"pending_admin_review", "approved", "rejected"}


def sanitize_proposal_text(value: Any, limit: int = 1600) -> str:
  return re.sub(r"\s+", " ", str(value or "").strip())[:limit]


def sanitize_institution_proposal_changes(changes: dict[str, Any] | None) -> dict[str, Any]:
  if not isinstance(changes, dict):
    return {}
  clean: dict[str, Any] = {}
  for field in INSTITUTION_PROPOSAL_FIELDS:
    if field not in changes:
      continue
    value = changes.get(field)
    if field == "careers":
      if isinstance(value, str):
        items = re.split(r"[\n,]+", value)
      elif isinstance(value, list):
        items = value
      else:
        items = []
      clean_items = [sanitize_proposal_text(item, 120) for item in items]
      clean_items = [item for item in clean_items if item][:20]
      if clean_items:
        clean[field] = clean_items
      continue
    limit = 2600 if field in {"requirementsSummary", "overview", "sourceNote", "feeNote"} else 600
    text = sanitize_proposal_text(value, limit)
    if text:
      clean[field] = text
  return clean


def normalize_institution_proposal(proposal: dict[str, Any]) -> dict[str, Any] | None:
  if not isinstance(proposal, dict):
    return None
  proposal_id = sanitize_proposal_text(proposal.get("id"), 80) or f"iprop-{uuid.uuid4().hex[:14]}"
  programme_id = sanitize_proposal_text(proposal.get("programmeId") or proposal.get("programme_id"), 140)
  institution = normalize_institution_name(proposal.get("institution"))
  changes = sanitize_institution_proposal_changes(proposal.get("changes"))
  if not programme_id or not institution or not changes:
    return None
  status = sanitize_proposal_text(proposal.get("status"), 80) or "pending_admin_review"
  if status not in INSTITUTION_PROPOSAL_STATUSES:
    status = "pending_admin_review"
  return {
    "id": proposal_id,
    "programmeId": programme_id,
    "programmeName": sanitize_proposal_text(proposal.get("programmeName"), 220),
    "institution": institution,
    "changes": changes,
    "note": sanitize_proposal_text(proposal.get("note"), 1200),
    "status": status,
    "requestedBy": proposal.get("requestedBy") if isinstance(proposal.get("requestedBy"), dict) else {},
    "reviewedBy": proposal.get("reviewedBy") if isinstance(proposal.get("reviewedBy"), dict) else {},
    "reviewNote": sanitize_proposal_text(proposal.get("reviewNote"), 1200),
    "createdAt": sanitize_proposal_text(proposal.get("createdAt"), 80) or now_iso(),
    "updatedAt": sanitize_proposal_text(proposal.get("updatedAt"), 80) or now_iso(),
  }


def load_institution_proposals() -> list[dict[str, Any]]:
  payload = load_state_payload(INSTITUTION_PROPOSAL_STATE_KEY, {"proposals": []})
  raw_proposals = payload.get("proposals", []) if isinstance(payload, dict) else payload if isinstance(payload, list) else []
  proposals = []
  for item in raw_proposals:
    proposal = normalize_institution_proposal(item)
    if proposal:
      proposals.append(proposal)
  return sorted(proposals, key=lambda item: item.get("createdAt") or "", reverse=True)


def save_institution_proposals(proposals: list[dict[str, Any]]) -> None:
  save_state_payload(
    INSTITUTION_PROPOSAL_STATE_KEY,
    {
      "proposals": proposals,
      "savedAt": now_iso(),
    },
  )


def get_visible_institution_proposals(actor: dict[str, Any], proposals: list[dict[str, Any]]) -> list[dict[str, Any]]:
  if actor.get("role") in {"owner", "admin"}:
    return proposals
  managed_institution = get_user_managed_institution(actor)
  return [proposal for proposal in proposals if institutions_match(proposal.get("institution"), managed_institution)]


AUDIT_EVENT_PREFIXES = ("admin_", "auth_", "registration_", "password_reset_", "institution_")
AUDIT_EVENT_TYPES = {
  "document_uploaded",
  "document_ocr_requested",
  "document_ocr_failed",
  "ai_guidance",
  "ai_compare",
  "course_search",
  "course_profile_viewed",
  "school_profile_viewed",
}


def is_audit_event(event_type: str | None) -> bool:
  clean_type = sanitize_event_type(event_type)
  return bool(clean_type) and (clean_type.startswith(AUDIT_EVENT_PREFIXES) or clean_type in AUDIT_EVENT_TYPES)


def format_event_label(event_type: str | None) -> str:
  clean_type = sanitize_event_type(event_type)
  return re.sub(r"[_:.-]+", " ", clean_type).strip().title() or "System event"


def build_admin_audit_log(limit: int = 120) -> list[dict[str, Any]]:
  capped_limit = max(10, min(limit, 300))
  users_by_id = {user["id"]: public_user(user) for user in get_auth_users_internal()}
  rows: list[dict[str, Any]] = []
  for event in safe_list_runtime_events(1200):
    event_type = event.get("event_type")
    if not is_audit_event(event_type):
      continue
    actor = users_by_id.get(str(event.get("user_id")), {"id": event.get("user_id"), "name": "System / anonymous"})
    payload = sanitize_event_payload(event.get("payload") if isinstance(event.get("payload"), dict) else parse_jsonish(event.get("payload"), {}))
    rows.append(
      {
        "id": event.get("id"),
        "eventType": sanitize_event_type(event_type),
        "label": event.get("label") or format_event_label(event_type),
        "actor": {
          "id": actor.get("id"),
          "name": actor.get("name") or "System / anonymous",
          "email": mask_email(actor.get("email")),
          "role": actor.get("role") or "system",
        },
        "payload": payload if isinstance(payload, dict) else {},
        "createdAt": event.get("created_at"),
      }
    )
    if len(rows) >= capped_limit:
      break
  return rows


def verification_payload_from_request(payload: AuthRegisterRequest) -> dict[str, Any]:
  password_salt, password_hash = hash_password(payload.password or "")
  return {
    "name": (payload.name or "").strip(),
    "email": normalize_email(payload.email),
    "district": (payload.district or "").strip(),
    "passwordSalt": password_salt,
    "passwordHash": password_hash,
  }


def get_latest_email_verification(email: str, purpose: str = "registration") -> dict[str, Any] | None:
  if using_supabase():
    rows = supabase_request(
      "GET",
      supabase_table_path(
        "runtime_email_verifications",
        f"email=eq.{supabase_filter_value(email)}&purpose=eq.{supabase_filter_value(purpose)}&consumed_at=is.null&order=created_at.desc&limit=1",
      ),
    )
    return rows[0] if rows else None

  with get_db_connection() as connection:
    row = connection.execute(
      """
      select * from email_verifications
      where email = ? and purpose = ? and consumed_at is null
      order by created_at desc
      limit 1
      """,
      (email, purpose),
    ).fetchone()
  return dict(row) if row else None


def consume_existing_email_verifications(email: str, purpose: str = "registration") -> None:
  timestamp = now_iso()
  if using_supabase():
    supabase_request(
      "PATCH",
      supabase_table_path(
        "runtime_email_verifications",
        f"email=eq.{supabase_filter_value(email)}&purpose=eq.{supabase_filter_value(purpose)}&consumed_at=is.null",
      ),
      {"consumed_at": timestamp},
      prefer="return=minimal",
    )
    return

  with get_db_connection() as connection:
    connection.execute(
      "update email_verifications set consumed_at = ? where email = ? and purpose = ? and consumed_at is null",
      (timestamp, email, purpose),
    )
    connection.commit()


def create_email_verification(email: str, payload: dict[str, Any], code: str, purpose: str = "registration") -> dict[str, Any]:
  latest = get_latest_email_verification(email, purpose)
  latest_created = parse_iso_datetime(str(latest.get("created_at") or "")) if latest else None
  if latest_created and datetime.now(timezone.utc) - latest_created < timedelta(seconds=EMAIL_VERIFICATION_RESEND_SECONDS):
    wait = EMAIL_VERIFICATION_RESEND_SECONDS - int((datetime.now(timezone.utc) - latest_created).total_seconds())
    raise HTTPException(status_code=429, detail=f"Please wait {max(wait, 1)} seconds before requesting another code.")

  consume_existing_email_verifications(email, purpose)
  code_salt, code_hash = hash_password(code)
  record = {
    "id": f"verify-{uuid.uuid4().hex[:16]}",
    "email": email,
    "purpose": purpose,
    "code_hash": code_hash,
    "code_salt": code_salt,
    "payload": payload,
    "attempts": 0,
    "created_at": now_iso(),
    "expires_at": minutes_from_now(EMAIL_VERIFICATION_TTL_MINUTES),
    "consumed_at": None,
  }
  if using_supabase():
    supabase_request(
      "POST",
      supabase_table_path("runtime_email_verifications"),
      record,
      prefer="return=minimal",
    )
    return record

  with get_db_connection() as connection:
    connection.execute(
      """
      insert into email_verifications
      (id, email, purpose, code_hash, code_salt, payload, attempts, created_at, expires_at, consumed_at)
      values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
      """,
      (
        record["id"],
        record["email"],
        record["purpose"],
        record["code_hash"],
        record["code_salt"],
        json.dumps(record["payload"], ensure_ascii=False),
        record["attempts"],
        record["created_at"],
        record["expires_at"],
        record["consumed_at"],
      ),
    )
    connection.commit()
  return record


def update_email_verification(record_id: str, changes: dict[str, Any]) -> None:
  if using_supabase():
    supabase_request(
      "PATCH",
      supabase_table_path("runtime_email_verifications", f"id=eq.{supabase_filter_value(record_id)}"),
      changes,
      prefer="return=minimal",
    )
    return

  allowed = {key: value for key, value in changes.items() if key in {"attempts", "consumed_at"}}
  if not allowed:
    return
  assignments = ", ".join(f"{key} = ?" for key in allowed)
  values = list(allowed.values()) + [record_id]
  with get_db_connection() as connection:
    connection.execute(f"update email_verifications set {assignments} where id = ?", values)
    connection.commit()


def get_verification_payload(record: dict[str, Any]) -> dict[str, Any]:
  payload = record.get("payload")
  if isinstance(payload, dict):
    return payload
  if isinstance(payload, str):
    return parse_jsonish(payload, {})
  return {}


def verify_email_code(email: str, code: str, purpose: str) -> dict[str, Any]:
  record = get_latest_email_verification(email, purpose)
  if not record:
    raise HTTPException(status_code=400, detail="No active verification code was found. Request a new code.")
  expires_at = parse_iso_datetime(str(record.get("expires_at") or ""))
  if not expires_at or expires_at < datetime.now(timezone.utc):
    update_email_verification(record["id"], {"consumed_at": now_iso()})
    raise HTTPException(status_code=400, detail="Verification code expired. Request a new code.")
  attempts = int(record.get("attempts") or 0)
  if attempts >= EMAIL_VERIFICATION_MAX_ATTEMPTS:
    update_email_verification(record["id"], {"consumed_at": now_iso()})
    raise HTTPException(status_code=400, detail="Too many incorrect attempts. Request a new code.")
  attempts += 1
  update_email_verification(record["id"], {"attempts": attempts})
  if not verify_password(code.strip(), {"passwordHash": record.get("code_hash"), "passwordSalt": record.get("code_salt")}):
    raise HTTPException(status_code=400, detail="Verification code is not correct.")
  update_email_verification(record["id"], {"consumed_at": now_iso()})
  return get_verification_payload(record)


def verify_registration_code(email: str, code: str) -> dict[str, Any]:
  return verify_email_code(email, code, "registration")


def validate_registration_input(payload: AuthRegisterRequest) -> tuple[str, str, str, str]:
  name = (payload.name or "").strip()
  email = normalize_email(payload.email)
  password = payload.password or ""
  district = (payload.district or "").strip()
  if not name or not email or not password:
    raise HTTPException(status_code=400, detail="Name, email, and password are required.")
  if len(password) < 8:
    raise HTTPException(status_code=400, detail="Password must be at least 8 characters.")
  return name, email, password, district


def create_student_account(registration: dict[str, Any], email_verified_at: str | None = None) -> dict[str, Any]:
  name = (registration.get("name") or "").strip()
  email = normalize_email(registration.get("email"))
  password_salt = registration.get("passwordSalt")
  password_hash = registration.get("passwordHash")
  if not name or not email or not password_salt or not password_hash:
    raise HTTPException(status_code=400, detail="Verification record is incomplete. Request a new code.")
  users = get_auth_users_internal()
  if any(user["email"] == email for user in users):
    raise HTTPException(status_code=409, detail="That email already has an account.")
  timestamp = now_iso()
  user = {
    "id": f"user-{uuid.uuid4().hex[:12]}",
    "name": name,
    "email": email,
    "passwordSalt": password_salt,
    "passwordHash": password_hash,
    "role": "student",
    "status": "active",
    "district": (registration.get("district") or "").strip(),
    "stream": "",
    "leavingYear": "",
    "incomeBand": "mid",
    "needSignals": [],
    "preferenceText": "",
    "grades": {},
    "documents": [],
    "shortlist": [],
    "createdAt": timestamp,
    "emailVerifiedAt": email_verified_at or timestamp,
    "reviewedAt": None,
    "lastActiveAt": timestamp,
    "lastActivity": "Student account created",
    "activity": [],
  }
  add_user_activity(user, "account_created", "Student account created after email verification", user)
  users.append(user)
  save_auth_users_internal(users)
  return user


def seed_bootstrap_admin() -> None:
  admin_email = normalize_email(os.getenv("ADMIN_EMAIL"))
  admin_password = (os.getenv("ADMIN_PASSWORD") or "").strip()
  if not admin_email or not admin_password:
    return
  users = get_auth_users_internal()
  existing = next((user for user in users if user["email"] == admin_email), None)
  salt, password_hash = hash_password(admin_password)
  timestamp = now_iso()
  admin_payload = {
    "name": (os.getenv("ADMIN_NAME") or "System Admin").strip(),
    "email": admin_email,
    "passwordSalt": salt,
    "passwordHash": password_hash,
    "role": "owner",
    "status": "active",
    "district": (os.getenv("ADMIN_DISTRICT") or "").strip(),
    "phone": (os.getenv("ADMIN_PHONE") or "").strip(),
    "emailVerifiedAt": timestamp,
    "reviewedAt": timestamp,
  }
  if existing:
    existing.update(admin_payload)
    existing.pop("password", None)
    add_user_activity(existing, "admin_bootstrap", "System admin synced from server environment", existing)
  else:
    admin_user = {
      "id": "admin-owner",
      **admin_payload,
      "stream": "",
      "leavingYear": "",
      "incomeBand": "mid",
      "needSignals": [],
      "preferenceText": "",
      "grades": {},
      "documents": [],
      "shortlist": [],
      "createdAt": timestamp,
      "emailVerifiedAt": timestamp,
      "lastActiveAt": timestamp,
      "lastActivity": "System admin created from server environment",
      "activity": [],
    }
    add_user_activity(admin_user, "admin_bootstrap", "System admin created from server environment", admin_user)
    users.insert(0, admin_user)
  for user in users:
    if user["email"] != admin_email and user.get("role") == "owner":
      user["role"] = "student"
      user["reviewedAt"] = None
      add_user_activity(
        user,
        "owner_demoted",
        "Owner access removed because system ownership is controlled by server environment",
        existing or user,
      )
  save_auth_users_internal(users)


def hash_session_token(token: str) -> str:
  return hashlib.sha256((token or "").encode("utf-8")).hexdigest()


def normalize_session_token(token: str | None) -> str:
  clean = (token or "").strip()
  if not clean:
    raise HTTPException(status_code=401, detail="Authentication required.")
  if len(clean) < AUTH_SESSION_TOKEN_MIN_LENGTH or len(clean) > AUTH_SESSION_TOKEN_MAX_LENGTH:
    raise HTTPException(status_code=401, detail="Authentication required.")
  if not re.fullmatch(r"[A-Za-z0-9._~\-+=:/]+", clean):
    raise HTTPException(status_code=401, detail="Authentication required.")
  return clean


def get_bearer_token(authorization: str | None) -> str:
  if not authorization or not authorization.lower().startswith("bearer "):
    raise HTTPException(status_code=401, detail="Authentication required.")
  token = normalize_session_token(authorization.split(" ", 1)[1].strip())
  return token


def get_session_user_id(token: str) -> str | None:
  normalized_token = normalize_session_token(token)
  token_hash = hash_session_token(normalized_token)
  if using_supabase():
    rows = supabase_request(
      "GET",
      supabase_table_path(
        "runtime_auth_sessions",
        f"token=eq.{supabase_filter_value(token_hash)}&select=user_id,created_at,last_seen_at&limit=1",
      ),
    )
    if not rows:
      legacy_rows = supabase_request(
        "GET",
        supabase_table_path(
          "runtime_auth_sessions",
          f"token=eq.{supabase_filter_value(normalized_token)}&select=user_id,created_at,last_seen_at&limit=1",
        ),
      )
      if not legacy_rows:
        return None
      rows = legacy_rows
    row = rows[0]
    if session_expired(str(row.get("last_seen_at") or ""), str(row.get("created_at") or "")):
      delete_auth_session(normalized_token)
      return None
    supabase_request(
      "PATCH",
      supabase_table_path("runtime_auth_sessions", f"token=eq.{supabase_filter_value(token_hash if rows[0].get('token') != normalized_token else normalized_token)}"),
      {"last_seen_at": now_iso()},
      prefer="return=minimal",
    )
    return row.get("user_id")

  with get_db_connection() as connection:
    row = connection.execute(
      "select user_id, created_at, last_seen_at, token from auth_sessions where token = ? or token = ?",
      (token_hash, normalized_token),
    ).fetchone()
    if row:
      if session_expired(row["last_seen_at"], row["created_at"]):
        connection.execute("delete from auth_sessions where token = ? or token = ?", (row["token"], normalized_token))
        connection.commit()
        return None
      connection.execute(
        "update auth_sessions set last_seen_at = current_timestamp where token = ? or token = ?",
        (row["token"], normalized_token),
      )
      connection.commit()
  return row["user_id"] if row else None

def require_current_user(authorization: str | None) -> dict[str, Any]:
  token = get_bearer_token(authorization)
  user_id = get_session_user_id(token)
  if not user_id:
    raise HTTPException(status_code=401, detail="Session expired. Please login again.")
  user = next((item for item in get_auth_users_internal() if item["id"] == user_id), None)
  if not user or user.get("status") == "suspended":
    raise HTTPException(status_code=403, detail="Account is not active.")
  return user


def require_admin_user(authorization: str | None) -> dict[str, Any]:
  user = require_current_user(authorization)
  if user.get("role") not in {"owner", "admin"}:
    raise HTTPException(status_code=403, detail="Admin access required.")
  return user


def require_institution_user(authorization: str | None) -> dict[str, Any]:
  user = require_current_user(authorization)
  if user.get("role") not in {"owner", "admin", "institution_admin"}:
    raise HTTPException(status_code=403, detail="Institution access required.")
  return user


def require_user_access(target_user_id: str, authorization: str | None) -> dict[str, Any]:
  user = require_current_user(authorization)
  if user.get("role") in {"owner", "admin"} or user.get("id") == target_user_id:
    return user
  raise HTTPException(status_code=403, detail="You can only access your own records.")


def create_auth_session(user_id: str) -> str:
  token = secrets.token_urlsafe(32)
  stored_token = hash_session_token(token)
  if using_supabase():
    supabase_request(
      "POST",
      supabase_table_path("runtime_auth_sessions"),
      {"token": stored_token, "user_id": user_id},
      prefer="return=minimal",
    )
    return token

  with get_db_connection() as connection:
    connection.execute("insert into auth_sessions (token, user_id) values (?, ?)", (stored_token, user_id))
    connection.commit()
  return token


def delete_auth_session(token: str) -> None:
  normalized_token = normalize_session_token(token)
  stored_token = hash_session_token(normalized_token)
  if using_supabase():
    supabase_request(
      "DELETE",
      supabase_table_path("runtime_auth_sessions", f"token=eq.{supabase_filter_value(stored_token)}"),
      prefer="return=minimal",
    )
    supabase_request(
      "DELETE",
      supabase_table_path("runtime_auth_sessions", f"token=eq.{supabase_filter_value(normalized_token)}"),
      prefer="return=minimal",
    )
    return

  with get_db_connection() as connection:
    connection.execute("delete from auth_sessions where token = ? or token = ?", (stored_token, normalized_token))
    connection.commit()


def delete_all_auth_sessions_for_user(user_id: str) -> None:
  """Delete all active sessions for a user. Used after password reset to invalidate stolen tokens."""
  if using_supabase():
    supabase_request(
      "DELETE",
      supabase_table_path("runtime_auth_sessions", f"user_id=eq.{supabase_filter_value(user_id)}"),
      prefer="return=minimal",
    )
    return

  with get_db_connection() as connection:
    connection.execute("delete from auth_sessions where user_id = ?", (user_id,))
    connection.commit()


def sanitize_database_state_payload(state_key: str, payload: Any) -> Any:
  if state_key != "auth_users" or not isinstance(payload, dict):
    return payload
  users = payload.get("users")
  if not isinstance(users, list):
    return payload
  cleaned_users = [public_user(user) for user in get_auth_users_internal()]
  return {**payload, "users": cleaned_users}


STARTUP_PERSISTENCE_ERROR = None
try:
  seed_bootstrap_admin()
  cleanup_expired_security_records(True)
except Exception as exc:
  STARTUP_PERSISTENCE_ERROR = str(exc)


@app.get("/api/db/state")
def get_database_state(request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  current_user = require_current_user(authorization)
  check_rate_limit(request, "db_state", 60, 60, current_user["id"])

  state = {}
  updated_at = {}

  # Only expose shared state that authenticated EduGuide users actually need.
  allowed_state_keys = {"review_state"}

  for row in list_state_payloads():
    state_key = row.get("state_key")

    if state_key not in allowed_state_keys:
      continue

    state[state_key] = parse_jsonish(row.get("payload"), row.get("payload"))
    updated_at[state_key] = row.get("updated_at")

  return {
    "ok": True,
    "database": get_data_backend(),
    "state": state,
    "updated_at": updated_at,
  }

@app.put("/api/db/state/{state_key}")
async def put_database_state(state_key: str, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  require_admin_user(authorization)
  safe_key = re.sub(r"[^a-zA-Z0-9_-]", "", state_key).strip()
  if not safe_key:
    return {"ok": False, "error": "Invalid state key"}
  if safe_key == "auth_users":
    return {"ok": True, "state_key": safe_key, "ignored": True}
  try:
    body = await request.json()
  except Exception as exc:
    raise HTTPException(status_code=400, detail="Request body must be JSON.") from exc
  payload = body.get("payload", body) if isinstance(body, dict) else body
  payload = sanitize_database_state_payload(safe_key, payload)
  serialized = json.dumps(payload, ensure_ascii=False, default=str)
  if len(serialized.encode("utf-8")) > 250000:
    raise HTTPException(status_code=413, detail="State payload exceeds the supported size limit.")
  save_state_payload(safe_key, payload)
  return {"ok": True, "state_key": safe_key}


@app.delete("/api/db/state/{state_key}")
def delete_database_state(state_key: str, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  require_admin_user(authorization)
  safe_key = re.sub(r"[^a-zA-Z0-9_-]", "", state_key).strip()
  if not safe_key:
    return {"ok": False, "error": "Invalid state key"}
  deleted = delete_state_payload(safe_key)
  return {"ok": True, "state_key": safe_key, "deleted": deleted}


@app.get("/api/auth/bootstrap")
def auth_bootstrap() -> dict[str, Any]:
  maybe_cleanup_security_records()
  seed_bootstrap_admin()
  admin_ready = any(user.get("role") in {"owner", "admin"} for user in get_auth_users_internal())
  return {"ok": True, "adminReady": admin_ready}


@app.post("/api/auth/login")
def auth_login(payload: AuthLoginRequest, request: Request) -> dict[str, Any]:
  maybe_cleanup_security_records()
  check_rate_limit(request, "auth_login", 8, 600, normalize_email(payload.email))
  seed_bootstrap_admin()
  email = normalize_email(payload.email)
  user = next((item for item in get_auth_users_internal() if item["email"] == email), None)
  if not user or not verify_password(payload.password, user):
    safe_insert_runtime_event("anonymous", "auth_login_failed", "Failed login attempt", {"email": mask_email(email), "ip": get_request_ip(request)})
    raise HTTPException(status_code=401, detail="Email or password is not correct.")
  if user.get("status") == "suspended":
    safe_insert_runtime_event(user["id"], "auth_login_blocked", "Suspended account login blocked", {"ip": get_request_ip(request)})
    raise HTTPException(status_code=403, detail="This account is suspended.")
  users = get_auth_users_internal()
  stored = next((item for item in users if item["id"] == user["id"]), user)
  if stored.get("password"):
    salt, password_hash = hash_password(payload.password)
    stored["passwordSalt"] = salt
    stored["passwordHash"] = password_hash
    stored.pop("password", None)
  add_user_activity(stored, "login", "Logged in", stored)
  save_auth_users_internal(users)
  safe_insert_runtime_event(stored["id"], "auth_login_success", "Logged in", {"role": stored.get("role"), "ip": get_request_ip(request)})
  return {"ok": True, "token": create_auth_session(stored["id"]), "user": public_user(stored)}


@app.post("/api/auth/register/request-code")
def auth_register_request_code(payload: AuthRegisterRequest, request: Request) -> dict[str, Any]:
  maybe_cleanup_security_records()
  check_rate_limit(request, "registration_code", 4, 900, normalize_email(payload.email))
  seed_bootstrap_admin()
  name, email, _password, _district = validate_registration_input(payload)
  if any(user["email"] == email for user in get_auth_users_internal()):
    raise HTTPException(status_code=409, detail="That email already has an account.")

  code = f"{secrets.randbelow(1_000_000):06d}"
  record = create_email_verification(email, verification_payload_from_request(payload), code)
  try:
    send_verification_email(email, name, code)
    safe_insert_runtime_event("anonymous", "registration_code_requested", "Registration email code sent", {"email": mask_email(email), "district": _district, "ip": get_request_ip(request)})
    return {
      "ok": True,
      "email": email,
      "emailSent": True,
      "expiresInMinutes": EMAIL_VERIFICATION_TTL_MINUTES,
      "resendSeconds": EMAIL_VERIFICATION_RESEND_SECONDS,
    }
  except Exception as exc:
    if email_debug_codes_enabled():
      return {
        "ok": True,
        "email": email,
        "emailSent": False,
        "debugCode": code,
        "expiresInMinutes": EMAIL_VERIFICATION_TTL_MINUTES,
        "resendSeconds": EMAIL_VERIFICATION_RESEND_SECONDS,
        "message": f"Email delivery is not configured. Development code: {code}",
      }
    update_email_verification(record["id"], {"consumed_at": now_iso()})
    missing = smtp_missing_keys()
    if missing:
      detail = f"Email delivery is not fully configured. Missing: {', '.join(missing)}."
    else:
      detail = f"Email delivery failed. SMTP is present, but the provider rejected the send or timed out. ({exc})"
    raise HTTPException(
      status_code=503,
      detail=detail,
    )


@app.post("/api/auth/register/verify")
def auth_register_verify(payload: AuthVerifyRegistrationRequest, request: Request) -> dict[str, Any]:
  maybe_cleanup_security_records()
  check_rate_limit(request, "registration_verify", 10, 900, normalize_email(payload.email))
  seed_bootstrap_admin()
  name, email, password, district = validate_registration_input(payload)
  if any(user["email"] == email for user in get_auth_users_internal()):
    raise HTTPException(status_code=409, detail="That email already has an account.")
  registration = verify_registration_code(email, payload.code)
  if normalize_email(registration.get("email")) != email:
    raise HTTPException(status_code=400, detail="Verification email does not match this registration.")
  if (registration.get("name") or "").strip() != name or (registration.get("district") or "").strip() != district:
    raise HTTPException(status_code=400, detail="Registration details changed. Request a new code.")
  if not verify_password(password, registration):
    raise HTTPException(status_code=400, detail="Password changed. Request a new code.")
  user = create_student_account(registration, email_verified_at=now_iso())
  safe_insert_runtime_event(user["id"], "registration_completed", "Student account verified", {"district": user.get("district"), "ip": get_request_ip(request)})
  return {"ok": True, "token": create_auth_session(user["id"]), "user": public_user(user)}


@app.post("/api/auth/register")
def auth_register(payload: AuthRegisterRequest, request: Request) -> dict[str, Any]:
  if not payload.code:
    raise HTTPException(status_code=400, detail="Request and verify an email code before creating an account.")
  verify_payload = AuthVerifyRegistrationRequest(**payload.model_dump())
  return auth_register_verify(verify_payload, request)


@app.post("/api/auth/password-reset/request-code")
def auth_password_reset_request_code(payload: AuthPasswordResetRequest, request: Request) -> dict[str, Any]:
  maybe_cleanup_security_records()
  email = normalize_email(payload.email)
  check_rate_limit(request, "password_reset_code", 4, 900, email)
  if not email or "@" not in email:
    raise HTTPException(status_code=400, detail="Enter the email address on your account.")

  users = get_auth_users_internal()
  user = next((item for item in users if item["email"] == email), None)
  if not user:
    return {
      "ok": True,
      "email": email,
      "emailSent": False,
      "expiresInMinutes": EMAIL_VERIFICATION_TTL_MINUTES,
      "resendSeconds": EMAIL_VERIFICATION_RESEND_SECONDS,
      "message": "If that email exists, a reset code will be sent.",
    }

  code = f"{secrets.randbelow(1_000_000):06d}"
  record = create_email_verification(email, {"email": email}, code, "password_reset")
  try:
    send_password_reset_email(email, user.get("name") or "student", code)
    safe_insert_runtime_event(user["id"], "password_reset_requested", "Requested password reset code", {})
    return {
      "ok": True,
      "email": email,
      "emailSent": True,
      "expiresInMinutes": EMAIL_VERIFICATION_TTL_MINUTES,
      "resendSeconds": EMAIL_VERIFICATION_RESEND_SECONDS,
    }
  except Exception as exc:
    if email_debug_codes_enabled():
      return {
        "ok": True,
        "email": email,
        "emailSent": False,
        "debugCode": code,
        "expiresInMinutes": EMAIL_VERIFICATION_TTL_MINUTES,
        "resendSeconds": EMAIL_VERIFICATION_RESEND_SECONDS,
        "message": f"Email delivery is not configured. Development reset code: {code}",
      }
    update_email_verification(record["id"], {"consumed_at": now_iso()})
    missing = smtp_missing_keys()
    detail = (
      f"Email delivery is not fully configured. Missing: {', '.join(missing)}."
      if missing
      else f"Email delivery failed. The provider rejected the password reset email or timed out. ({exc})"
    )
    raise HTTPException(status_code=503, detail=detail)


@app.post("/api/auth/password-reset/confirm")
def auth_password_reset_confirm(payload: AuthPasswordResetConfirmRequest, request: Request) -> dict[str, Any]:
  maybe_cleanup_security_records()
  email = normalize_email(payload.email)
  check_rate_limit(request, "password_reset_confirm", 10, 900, email)
  if len(payload.password or "") < 8:
    raise HTTPException(status_code=400, detail="Password must be at least 8 characters.")

  registration = verify_email_code(email, payload.code, "password_reset")
  if normalize_email(registration.get("email")) != email:
    raise HTTPException(status_code=400, detail="Verification email does not match this reset request.")

  users = get_auth_users_internal()
  user = next((item for item in users if item["email"] == email), None)
  if not user:
    raise HTTPException(status_code=404, detail="Account not found.")

  salt, password_hash = hash_password(payload.password)
  user["passwordSalt"] = salt
  user["passwordHash"] = password_hash
  user.pop("password", None)
  user["status"] = user.get("status") or "active"
  add_user_activity(user, "password_reset", "Password reset by email verification", user)
  safe_insert_runtime_event(user["id"], "password_reset_confirmed", "Password reset completed", {})
  save_auth_users_internal(users)
  delete_all_auth_sessions_for_user(user["id"])
  return {"ok": True, "token": create_auth_session(user["id"]), "user": public_user(user)}


@app.get("/api/auth/me")
def auth_me(authorization: str | None = Header(default=None)) -> dict[str, Any]:
  user = require_current_user(authorization)
  return {"ok": True, "user": public_user(user)}


@app.post("/api/auth/logout")
def auth_logout(request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  token = get_bearer_token(authorization)
  user_id = get_session_user_id(token)
  delete_auth_session(token)
  if user_id:
    safe_insert_runtime_event(user_id, "auth_logout", "Logged out", {"ip": get_request_ip(request)})
  return {"ok": True}


@app.put("/api/auth/me")
def auth_update_me(payload: AuthProfileRequest, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  current = require_current_user(authorization)
  safe_payload = sanitize_profile_payload(payload)
  users = get_auth_users_internal()
  user = next((item for item in users if item["id"] == current["id"]), None)
  if not user:
    raise HTTPException(status_code=404, detail="User not found.")
  for field in ["name", "district", "stream", "leavingYear", "incomeBand", "preferenceText"]:
    value = getattr(safe_payload, field)
    if value is not None:
      user[field] = value
  user["needSignals"] = safe_payload.needSignals
  user["grades"] = safe_payload.grades
  user["documents"] = safe_payload.documents
  user["shortlist"] = safe_payload.shortlist
  user["shortlistPathways"] = safe_payload.shortlistPathways
  add_user_activity(user, "profile_updated", "Updated profile", user)
  save_auth_users_internal(users)
  return {"ok": True, "user": public_user(user)}


def admin_public_user_summary(user: dict[str, Any]) -> dict[str, Any]:
  return {
    "id": user.get("id"),
    "name": user.get("name"),
    "email": user.get("email"),
    "role": user.get("role"),
    "managedInstitution": get_user_managed_institution(user),
    "status": user.get("status"),
    "district": user.get("district"),
    "createdAt": user.get("createdAt"),
    "reviewedAt": user.get("reviewedAt"),
    "lastActiveAt": user.get("lastActiveAt"),
    "lastActivity": user.get("lastActivity"),
  }


def admin_grade_points(grade: Any) -> int:
  return {"A*": 8, "A": 7, "B": 6, "C": 5, "D": 4, "E": 3, "F": 2, "G": 1, "X": 0, "Z": 0}.get(str(grade or "").upper(), 0)


def admin_grade_meets(grade: Any, minimum: str = "D") -> bool:
  return admin_grade_points(grade) >= admin_grade_points(minimum)


def admin_student_science_status(user: dict[str, Any]) -> dict[str, Any]:
  grades = user.get("grades") if isinstance(user.get("grades"), dict) else {}
  science_codes = ["PSCI", "BIO", "PHY", "CHEM", "AGR"]
  entered_science = [code for code in science_codes if grades.get(code)]
  return {
    "hasMath": bool(grades.get("MATH")),
    "hasStrongMath": admin_grade_meets(grades.get("MATH")),
    "hasScience": bool(entered_science),
    "hasStrongScience": any(admin_grade_meets(grades.get(code)) for code in entered_science),
    "scienceSubjects": entered_science,
  }


def list_admin_document_rows() -> list[dict[str, Any]]:
  if using_supabase():
    return supabase_request(
      "GET",
      supabase_table_path(
        "runtime_uploaded_documents",
        "select=id,user_id,original_name,status,extraction_status,extraction_error,uploaded_at&order=uploaded_at.desc&limit=10000",
      ),
    ) or []

  with get_db_connection() as connection:
    rows = connection.execute(
      """
      select id, user_id, original_name, status, extraction_status, extraction_error, uploaded_at
      from uploaded_documents
      order by uploaded_at desc
      limit 10000
      """
    ).fetchall()
  return [dict(row) for row in rows]


def list_admin_recommendation_rows(limit: int = 250) -> list[dict[str, Any]]:
  capped_limit = max(1, min(limit, 1000))
  if using_supabase():
    return supabase_request(
      "GET",
      supabase_table_path(
        "runtime_recommendation_runs",
        f"select=mode,question,profile_name,payload,created_at&order=created_at.desc&limit={capped_limit}",
      ),
    ) or []

  with get_db_connection() as connection:
    rows = connection.execute(
      """
      select mode, question, profile_name, payload, created_at
      from recommendation_runs
      order by id desc
      limit ?
      """,
      (capped_limit,),
    ).fetchall()
  return [dict(row) for row in rows]


def add_admin_programme_signal(signals: dict[str, dict[str, Any]], programme: dict[str, Any]) -> None:
  programme_id = str(programme.get("programmeId") or programme.get("id") or "").strip()
  programme_name = str(programme.get("programmeName") or programme.get("programme") or programme.get("name") or "").strip()
  institution = str(programme.get("institution") or "").strip()
  if not programme_id and not programme_name:
    return
  key = programme_id or f"{institution}:{programme_name}".lower()
  current = signals.setdefault(
    key,
    {
      "programmeId": programme_id or None,
      "programmeName": programme_name or programme_id or "Programme",
      "institution": institution,
      "score": 0,
      "saved": 0,
      "viewed": 0,
      "aiMentions": 0,
    },
  )
  source = str(programme.get("source") or "aiMentions")
  weight = int(programme.get("weight") or 1)
  current["score"] += weight
  if source not in {"saved", "viewed", "aiMentions"}:
    source = "aiMentions"
  current[source] = int(current.get(source) or 0) + 1
  if institution and not current.get("institution"):
    current["institution"] = institution
  if programme_name and current.get("programmeName") == current.get("programmeId"):
    current["programmeName"] = programme_name


def build_admin_intelligence() -> dict[str, Any]:
  users = get_auth_users_internal()
  students = [user for user in users if user.get("role") not in {"owner", "admin", "institution_admin"}]
  student_ids = {str(user["id"]) for user in students}
  users_by_id = {user["id"]: user for user in users}
  signals: dict[str, dict[str, Any]] = {}
  active_ai_user_ids: set[str] = set()
  runtime_events = safe_list_runtime_events(1500)
  search_counts: dict[str, int] = {}
  school_counts: dict[str, int] = {}
  blocked_reason_counts: dict[str, int] = {}

  for user in students:
    for programme_id in user.get("shortlist") or []:
      add_admin_programme_signal(signals, {"programmeId": programme_id, "source": "saved", "weight": 3})
    if runtime_events:
      continue
    for activity in user.get("activity") or []:
      metadata = activity.get("metadata") if isinstance(activity.get("metadata"), dict) else {}
      activity_type = activity.get("type")
      if activity_type in {"ai_guidance", "ai_compare"}:
        active_ai_user_ids.add(user["id"])
      if activity_type == "shortlist_updated":
        add_admin_programme_signal(signals, {**metadata, "source": "saved", "weight": 2})
      if activity_type == "programme_viewed":
        add_admin_programme_signal(signals, {**metadata, "source": "viewed", "weight": 1})

  recent_questions = []
  for row in list_admin_recommendation_rows():
    payload = parse_jsonish(row.get("payload"), {})
    for match in (payload.get("recommendable_matches") or [])[:3]:
      add_admin_programme_signal(
        signals,
        {
          "programmeName": match.get("programme"),
          "institution": match.get("institution"),
          "source": "aiMentions",
          "weight": 1,
        },
      )
    question = str(row.get("question") or "").strip()
    if question:
      recent_questions.append(
        {
          "question": question[:180],
          "profileName": row.get("profile_name"),
          "mode": row.get("mode"),
          "createdAt": row.get("created_at"),
        }
      )

  for event in runtime_events:
    payload = parse_jsonish(event.get("payload"), {})
    if not isinstance(payload, dict):
      payload = {}
    event_type = str(event.get("event_type") or "")
    event_user_id = str(event.get("user_id") or "")
    if event_type in {"ai_guidance", "ai_compare", "ai_chat_cleared"} and event_user_id in student_ids:
      active_ai_user_ids.add(event_user_id)
    if event_type in {"shortlist_updated", "programme_viewed", "course_profile_viewed"}:
      add_admin_programme_signal(
        signals,
        {
          "programmeId": payload.get("programmeId"),
          "programmeName": payload.get("programmeName"),
          "institution": payload.get("institution"),
          "source": "saved" if event_type == "shortlist_updated" else "viewed",
          "weight": 2 if event_type == "shortlist_updated" else 1,
        },
      )
    if event_type == "course_search":
      query = str(payload.get("query") or "").strip().lower()
      if len(query) >= 3:
        search_counts[query] = search_counts.get(query, 0) + 1
    if event_type == "school_profile_viewed":
      institution = str(payload.get("institution") or "").strip()
      if institution:
        school_counts[institution] = school_counts.get(institution, 0) + 1
    if event_type == "matches_calculated":
      for reason in payload.get("blockedReasons") or []:
        reason_text = str(reason or "").strip()
        if reason_text:
          blocked_reason_counts[reason_text] = blocked_reason_counts.get(reason_text, 0) + 1

  blocked_by_math_science = []
  for user in students:
    status = admin_student_science_status(user)
    reasons = []
    if not status["hasStrongMath"]:
      reasons.append("Mathematics missing or below D")
    if not status["hasStrongScience"]:
      reasons.append("Science/Agriculture gate missing or below D")
    if reasons:
      blocked_by_math_science.append(
        {
          **admin_public_user_summary(user),
          "reason": ", ".join(reasons),
          "scienceSubjects": status["scienceSubjects"],
        }
      )

  ocr_failures = []
  for document in list_admin_document_rows():
    status = str(document.get("extraction_status") or "")
    error = str(document.get("extraction_error") or "")
    label = str(document.get("status") or "")
    if status != "failed" and not error and "ocr failed" not in label.lower():
      continue
    user = users_by_id.get(str(document.get("user_id") or ""))
    ocr_failures.append(
      {
        "documentId": document.get("id"),
        "documentName": document.get("original_name"),
        "status": label or status,
        "error": error,
        "uploadedAt": document.get("uploaded_at"),
        "user": admin_public_user_summary(user) if user else {"id": document.get("user_id"), "name": "Unknown user"},
      }
    )

  runtime_warnings = []
  if not supabase_configured():
    runtime_warnings.append("Hosted runtime is not connected to Supabase; free container data can reset.")
  if not smtp_configured():
    runtime_warnings.append("SMTP is not configured; public email verification cannot send codes.")
  if STARTUP_PERSISTENCE_ERROR:
    runtime_warnings.append(f"Startup persistence warning: {STARTUP_PERSISTENCE_ERROR}")
  smtp_missing = smtp_missing_keys()
  if smtp_missing:
    runtime_warnings.append(f"SMTP missing safe key(s): {', '.join(smtp_missing)}")

  return {
    "ok": True,
    "generatedAt": now_iso(),
    "database": get_data_backend(),
    "studentsCount": len(students),
    "activeAiUsers": len(active_ai_user_ids) or len({item.get("profileName") for item in recent_questions if item.get("profileName")}),
    "topProgrammes": sorted(signals.values(), key=lambda item: (-int(item.get("score") or 0), item.get("programmeName") or ""))[:8],
    "topSearches": [
      {"query": key, "count": value}
      for key, value in sorted(search_counts.items(), key=lambda item: (-item[1], item[0]))[:8]
    ],
    "topSchools": [
      {"institution": key, "count": value}
      for key, value in sorted(school_counts.items(), key=lambda item: (-item[1], item[0]))[:8]
    ],
    "blockedReasons": [
      {"reason": key, "count": value}
      for key, value in sorted(blocked_reason_counts.items(), key=lambda item: (-item[1], item[0]))[:8]
    ],
    "missingWarnings": runtime_warnings,
    "blockedByMathScience": blocked_by_math_science[:12],
    "ocrFailures": ocr_failures[:12],
    "newUsers": [admin_public_user_summary(user) for user in students if not user.get("reviewedAt")][:12],
    "recentQuestions": recent_questions[:8],
  }


@app.get("/api/institution/proposals")
def list_institution_proposals(request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  check_rate_limit(request, "institution_proposals_list", 120, 3600)
  actor = require_institution_user(authorization)
  proposals = load_institution_proposals()
  visible = get_visible_institution_proposals(actor, proposals)
  return {
    "ok": True,
    "database": get_data_backend(),
    "managedInstitution": get_user_managed_institution(actor),
    "proposals": visible,
  }


@app.post("/api/institution/proposals")
def create_institution_proposal(payload: InstitutionProposalRequest, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  check_rate_limit(request, "institution_proposals_create", 35, 3600)
  actor = require_institution_user(authorization)
  institution = normalize_institution_name(payload.institution)
  if not institution:
    raise HTTPException(status_code=400, detail="Institution is required.")
  if actor.get("role") == "institution_admin" and not institutions_match(institution, get_user_managed_institution(actor)):
    raise HTTPException(status_code=403, detail="You can only submit updates for your assigned institution.")

  changes = sanitize_institution_proposal_changes(payload.changes)
  if not changes:
    raise HTTPException(status_code=400, detail="At least one proposed change is required.")
  programme_id = sanitize_proposal_text(payload.programmeId, 140)
  if not programme_id:
    raise HTTPException(status_code=400, detail="Programme is required.")

  proposals = load_institution_proposals()
  proposal = {
    "id": f"iprop-{uuid.uuid4().hex[:14]}",
    "programmeId": programme_id,
    "programmeName": sanitize_proposal_text(payload.programmeName, 220),
    "institution": institution,
    "changes": changes,
    "note": sanitize_proposal_text(payload.note, 1200),
    "status": "pending_admin_review",
    "requestedBy": {
      "id": actor.get("id"),
      "name": actor.get("name"),
      "email": mask_email(actor.get("email")),
      "role": actor.get("role"),
    },
    "reviewedBy": {},
    "reviewNote": "",
    "createdAt": now_iso(),
    "updatedAt": now_iso(),
  }
  proposals.insert(0, proposal)
  save_institution_proposals(proposals)
  safe_insert_runtime_event(
    actor["id"],
    "institution_proposal_submitted",
    "Institution submitted catalogue update",
    {
      "proposalId": proposal["id"],
      "programmeId": proposal["programmeId"],
      "programmeName": proposal["programmeName"],
      "institution": proposal["institution"],
      "fields": list(changes.keys()),
      "ip": get_request_ip(request),
    },
  )
  return {"ok": True, "proposal": proposal, "proposals": get_visible_institution_proposals(actor, proposals)}


@app.put("/api/institution/proposals/{proposal_id}")
def decide_institution_proposal(proposal_id: str, payload: InstitutionProposalDecisionRequest, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  check_rate_limit(request, "institution_proposals_decide", 80, 3600)
  actor = require_admin_user(authorization)
  if payload.status not in {"approved", "rejected"}:
    raise HTTPException(status_code=400, detail="Proposal status must be approved or rejected.")
  proposals = load_institution_proposals()
  proposal = next((item for item in proposals if item.get("id") == proposal_id), None)
  if not proposal:
    raise HTTPException(status_code=404, detail="Proposal not found.")
  proposal["status"] = payload.status
  proposal["reviewNote"] = sanitize_proposal_text(payload.note, 1200)
  proposal["reviewedBy"] = {
    "id": actor.get("id"),
    "name": actor.get("name"),
    "email": mask_email(actor.get("email")),
    "role": actor.get("role"),
  }
  proposal["updatedAt"] = now_iso()
  save_institution_proposals(proposals)
  safe_insert_runtime_event(
    actor["id"],
    "institution_proposal_reviewed",
    "Admin reviewed institution update",
    {
      "proposalId": proposal.get("id"),
      "programmeId": proposal.get("programmeId"),
      "institution": proposal.get("institution"),
      "status": proposal.get("status"),
      "ip": get_request_ip(request),
    },
  )
  return {"ok": True, "proposal": proposal, "proposals": get_visible_institution_proposals(actor, proposals)}


@app.get("/api/admin/users")
def admin_list_users(request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  check_rate_limit(request, "admin_users", 120, 3600)
  require_admin_user(authorization)
  return {"ok": True, "users": [public_user(user) for user in get_auth_users_internal()]}


@app.get("/api/admin/intelligence")
def admin_intelligence(request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  check_rate_limit(request, "admin_intelligence", 120, 3600)
  require_admin_user(authorization)
  return build_admin_intelligence()


@app.get("/api/admin/audit")
def admin_audit(request: Request, authorization: str | None = Header(default=None), limit: int = 120) -> dict[str, Any]:
  check_rate_limit(request, "admin_audit", 120, 3600)
  actor = require_admin_user(authorization)
  rows = build_admin_audit_log(limit)
  safe_insert_runtime_event(actor["id"], "admin_audit_viewed", "Admin opened audit log", {"rows": len(rows), "ip": get_request_ip(request)})
  return {
    "ok": True,
    "database": get_data_backend(),
    "generatedAt": now_iso(),
    "events": rows,
  }


@app.post("/api/admin/test-email")
def admin_test_email(request: Request, payload: AdminTestEmailRequest | None = None, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  check_rate_limit(request, "admin_test_email", 5, 900)
  actor = require_admin_user(authorization)
  target_email = normalize_email((payload.email if payload else None) or actor.get("email"))
  if not target_email or "@" not in target_email:
    raise HTTPException(status_code=400, detail="A valid test email address is required.")
  if not smtp_configured():
    missing = ", ".join(smtp_missing_keys())
    raise HTTPException(status_code=503, detail=f"SMTP is not fully configured. Missing: {missing}.")

  sent_at = now_iso()
  try:
    send_plain_email(
      target_email,
      "EduGuide LS production email test",
      [
        f"Hello {actor.get('name') or 'admin'},",
        "",
        "This is a production SMTP test from EduGuide LS.",
        f"Sent at: {sent_at}",
        "",
        "If you received this message, the server can send account verification emails.",
        "",
        "EduGuide LS",
      ],
    )
  except Exception as exc:
    raise HTTPException(status_code=502, detail=f"SMTP test failed: {exc}")

  safe_insert_runtime_event(actor["id"], "admin_test_email", "Admin sent SMTP test email", {"sentTo": mask_email(target_email), "ip": get_request_ip(request)})
  return {
    "ok": True,
    "sentTo": mask_email(target_email),
    "sentAt": sent_at,
    "message": "Test email sent. Check the inbox and spam folder for the admin email.",
  }


@app.put("/api/admin/users/{user_id}/role")
def admin_set_user_role(user_id: str, payload: UserRoleUpdate, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  check_rate_limit(request, "admin_user_role", 40, 3600)
  actor = require_admin_user(authorization)
  if payload.role not in {"admin", "student", "institution_admin"}:
    raise HTTPException(status_code=400, detail="Invalid role.")
  managed_institution = normalize_institution_name(payload.institution or payload.managedInstitution)
  if payload.role == "institution_admin" and not managed_institution:
    raise HTTPException(status_code=400, detail="Choose an institution before granting institution access.")
  users = get_auth_users_internal()
  user = next((item for item in users if item["id"] == user_id), None)
  if not user:
    raise HTTPException(status_code=404, detail="User not found.")
  if user["id"] == actor["id"] or user.get("role") == "owner":
    raise HTTPException(status_code=400, detail="Protected account.")
  user["role"] = payload.role
  if payload.role == "institution_admin":
    user["managedInstitution"] = managed_institution
  else:
    user["managedInstitution"] = ""
  if payload.role in {"admin", "institution_admin"}:
    user["reviewedAt"] = user.get("reviewedAt") or now_iso()
  role_label = payload.role.replace("_", " ")
  add_user_activity(user, "role_changed", f"Role changed to {role_label}", actor, {"role": payload.role, "managedInstitution": managed_institution})
  safe_insert_runtime_event(actor["id"], "admin_role_changed", "Admin changed a user role", {"targetUserId": user["id"], "targetEmail": mask_email(user.get("email")), "role": payload.role, "managedInstitution": managed_institution, "ip": get_request_ip(request)})
  save_auth_users_internal(users)
  return {"ok": True, "user": public_user(user), "users": [public_user(item) for item in users]}


@app.put("/api/admin/users/{user_id}/status")
def admin_set_user_status(user_id: str, payload: UserStatusUpdate, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  check_rate_limit(request, "admin_user_status", 40, 3600)
  actor = require_admin_user(authorization)
  if payload.status not in {"active", "suspended"}:
    raise HTTPException(status_code=400, detail="Invalid status.")
  users = get_auth_users_internal()
  user = next((item for item in users if item["id"] == user_id), None)
  if not user:
    raise HTTPException(status_code=404, detail="User not found.")
  if user["id"] == actor["id"] or user.get("role") == "owner":
    raise HTTPException(status_code=400, detail="Protected account.")
  user["status"] = payload.status
  if payload.status == "active":
    user["reviewedAt"] = user.get("reviewedAt") or now_iso()
  add_user_activity(user, "status_changed", f"Account {payload.status}", actor, {"status": payload.status})
  safe_insert_runtime_event(actor["id"], "admin_status_changed", "Admin changed account status", {"targetUserId": user["id"], "targetEmail": mask_email(user.get("email")), "status": payload.status, "ip": get_request_ip(request)})
  save_auth_users_internal(users)
  return {"ok": True, "user": public_user(user), "users": [public_user(item) for item in users]}


@app.put("/api/admin/users/{user_id}/review")
def admin_review_user(user_id: str, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  check_rate_limit(request, "admin_user_review", 80, 3600)
  actor = require_admin_user(authorization)
  users = get_auth_users_internal()
  user = next((item for item in users if item["id"] == user_id), None)
  if not user:
    raise HTTPException(status_code=404, detail="User not found.")
  user["reviewedAt"] = now_iso()
  add_user_activity(user, "reviewed", "Account reviewed by admin", actor)
  safe_insert_runtime_event(actor["id"], "admin_user_reviewed", "Admin reviewed a student account", {"targetUserId": user["id"], "targetEmail": mask_email(user.get("email")), "ip": get_request_ip(request)})
  save_auth_users_internal(users)
  return {"ok": True, "user": public_user(user), "users": [public_user(item) for item in users]}


@app.post("/api/documents/upload")
async def upload_documents(request: Request, user_id: str = Form(...), files: list[UploadFile] = File(...), authorization: str | None = Header(default=None)) -> dict[str, Any]:
  require_user_access(user_id, authorization)
  check_rate_limit(request, "document_upload", 12, 3600, user_id)
  safe_user_id = sanitize_storage_segment(user_id, "anonymous")
  if not files:
    raise HTTPException(status_code=400, detail="No documents were uploaded.")
  if len(files) > MAX_FILES_PER_UPLOAD:
    raise HTTPException(status_code=400, detail=f"Maximum {MAX_FILES_PER_UPLOAD} files per upload allowed.")

  saved_documents = []

  for upload in files:
    original_name = sanitize_upload_filename(upload.filename or "document")
    extension = Path(original_name).suffix.lower()
    if extension not in ALLOWED_UPLOAD_EXTENSIONS:
      raise HTTPException(status_code=400, detail=f"File type '{extension}' is not allowed. Supported types: {', '.join(sorted(ALLOWED_UPLOAD_EXTENSIONS))}")
    
    # Validate MIME type
    if not validate_mime_type(upload.content_type, extension):
      raise HTTPException(status_code=400, detail=f"MIME type '{upload.content_type}' does not match file extension '{extension}'.")
    
    stored_name = f"{uuid.uuid4().hex}{extension}"
    content_type = upload.content_type
    chunks: list[bytes] = []
    size = 0

    try:
      while True:
        chunk = await upload.read(1024 * 1024)
        if not chunk:
          break
        size += len(chunk)
        if size > MAX_UPLOAD_BYTES:
          raise HTTPException(status_code=413, detail=f"{original_name} is larger than 10 MB.")
        chunks.append(chunk)
    finally:
      await upload.close()

    document_id = uuid.uuid4().hex
    content = b"".join(chunks)
    
    # Validate file signature (magic numbers)
    if not validate_file_signature(content, extension):
      raise HTTPException(status_code=400, detail=f"File signature validation failed for '{extension}'. The file may be corrupted or not actually a {extension} file.")
    
    if using_supabase():
      storage_path = f"{safe_user_id}/{stored_name}"
      upload_supabase_storage_object(storage_path, content, content_type)
    else:
      user_folder = UPLOAD_ROOT / safe_user_id
      user_folder.mkdir(parents=True, exist_ok=True)
      destination = user_folder / stored_name
      with destination.open("wb") as output:
        output.write(content)
      storage_path = destination.relative_to(ROOT).as_posix()

    insert_document_record(
      {
        "id": document_id,
        "user_id": user_id,
        "original_name": original_name,
        "stored_name": stored_name,
        "content_type": content_type,
        "size_bytes": size,
        "storage_path": storage_path,
        "status": "Uploaded - OCR pending",
        "extraction_status": "pending",
        "extracted_grades": [],
      }
    )
    row = extract_document(document_id)
    saved_documents.append(document_response(row))

  return {"ok": True, "documents": saved_documents}


@app.get("/api/documents/user/{user_id}")
def list_user_documents(user_id: str, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  require_user_access(user_id, authorization)
  rows = list_document_records(user_id)
  return {"ok": True, "documents": [document_response(row) for row in rows]}


@app.post("/api/documents/{document_id}/extract")
def rerun_document_extraction(document_id: str, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  safe_document_id = sanitize_storage_segment(document_id)
  row = get_document_record(safe_document_id)
  if not row:
    raise HTTPException(status_code=404, detail="Document not found.")
  require_user_access(str(row_get(row, "user_id")), authorization)
  check_rate_limit(request, "document_ocr", 20, 3600, str(row_get(row, "user_id")))
  row = extract_document(safe_document_id)
  return {"ok": True, "document": document_response(row)}


@app.get("/api/documents/{document_id}/download")
def download_document(document_id: str, authorization: str | None = Header(default=None)) -> Response:
  safe_document_id = sanitize_storage_segment(document_id)
  row = get_document_record(safe_document_id)
  if not row:
    raise HTTPException(status_code=404, detail="Document not found.")
  require_user_access(str(row_get(row, "user_id")), authorization)
  if using_supabase():
    content = download_supabase_storage_object(str(row_get(row, "storage_path")))
    filename = sanitize_upload_filename(str(row_get(row, "original_name") or "document"))
    return Response(
      content=content,
      media_type=row_get(row, "content_type") or "application/octet-stream",
      headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

  path = ROOT / str(row_get(row, "storage_path"))
  if not path.exists():
    raise HTTPException(status_code=404, detail="Stored document file is missing.")
  return FileResponse(path, media_type=row_get(row, "content_type") or "application/octet-stream", filename=row_get(row, "original_name"))


@app.delete("/api/documents/{document_id}")
def delete_document(document_id: str, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  safe_document_id = sanitize_storage_segment(document_id)
  existing = get_document_record(safe_document_id)
  if existing:
    require_user_access(str(row_get(existing, "user_id")), authorization)
  row = delete_document_record(safe_document_id)
  if not row:
    return {"ok": True, "deleted": 0}
  if using_supabase():
    delete_supabase_storage_object(str(row_get(row, "storage_path")))
    return {"ok": True, "deleted": 1}

  path = ROOT / str(row_get(row, "storage_path"))
  if path.exists() and UPLOAD_ROOT in path.resolve().parents:
    path.unlink(missing_ok=True)
  return {"ok": True, "deleted": 1}


@app.get("/api/db/diagnostics")
def database_diagnostics(authorization: str | None = Header(default=None)) -> dict[str, Any]:
  require_admin_user(authorization)
  ai_provider = get_ai_provider()
  ai_model = get_gemini_model_candidates()[0] if ai_provider == "gemini" else (os.getenv("OPENAI_MODEL", "gpt-4o").strip() or "gpt-4o")
  if using_supabase():
    state_rows = list_state_payloads()
    document_rows = supabase_request(
      "GET",
      supabase_table_path("runtime_uploaded_documents", "select=id,user_id,extraction_status&limit=10000"),
    ) or []
    run_rows = supabase_request(
      "GET",
      supabase_table_path("runtime_recommendation_runs", "select=mode,question,profile_name,created_at&order=created_at.desc&limit=5"),
    ) or []
    run_count_rows = supabase_request(
      "GET",
      supabase_table_path("runtime_recommendation_runs", "select=id&limit=10000"),
    ) or []
    try:
      chat_count_rows = supabase_request(
        "GET",
        supabase_table_path("runtime_ai_chat_messages", "select=id&limit=10000"),
      ) or []
    except Exception:
      chat_count_rows = []
    try:
      event_count_rows = supabase_request(
        "GET",
        supabase_table_path("runtime_events", "select=id&limit=10000"),
      ) or []
    except Exception:
      event_count_rows = []
    user_counts: dict[str, int] = {}
    status_counts: dict[str, int] = {}
    for document in document_rows:
      user_counts[str(document.get("user_id") or "unknown")] = user_counts.get(str(document.get("user_id") or "unknown"), 0) + 1
      status = str(document.get("extraction_status") or "pending")
      status_counts[status] = status_counts.get(status, 0) + 1
    document_users = [{"user_id": key, "total": value} for key, value in sorted(user_counts.items(), key=lambda item: (-item[1], item[0]))[:10]]
    extraction_statuses = [{"extraction_status": key, "total": value} for key, value in sorted(status_counts.items(), key=lambda item: (-item[1], item[0]))]
    run_count = len(run_count_rows)
    chat_message_count = len(chat_count_rows)
    event_count = len(event_count_rows)
    document_count = len(document_rows)
    latest_runs = run_rows
  else:
    with get_db_connection() as connection:
      state_rows = [
        {
          "state_key": row["state_key"],
          "payload": parse_jsonish(row["payload"], {}),
          "updated_at": row["updated_at"],
        }
        for row in connection.execute("select state_key, payload, updated_at from app_state").fetchall()
      ]
      run_count = connection.execute("select count(*) as total from recommendation_runs").fetchone()["total"]
      chat_message_count = connection.execute("select count(*) as total from ai_chat_messages").fetchone()["total"]
      event_count = connection.execute("select count(*) as total from runtime_events").fetchone()["total"]
      document_count = connection.execute("select count(*) as total from uploaded_documents").fetchone()["total"]
      document_users = [
        dict(row)
        for row in connection.execute(
          """
          select user_id, count(*) as total
          from uploaded_documents
          group by user_id
          order by total desc, user_id
          limit 10
          """
        ).fetchall()
      ]
      extraction_statuses = [
        dict(row)
        for row in connection.execute(
          """
          select extraction_status, count(*) as total
          from uploaded_documents
          group by extraction_status
          order by total desc, extraction_status
          """
        ).fetchall()
      ]
      latest_runs = [
        dict(row)
        for row in connection.execute(
          """
          select mode, question, profile_name, created_at
          from recommendation_runs
          order by id desc
          limit 5
          """
        ).fetchall()
      ]
  state_counts = {}
  for row in state_rows:
    payload = parse_jsonish(row.get("payload"), {})
    state_key = row.get("state_key")
    if state_key == "auth_users":
      state_counts[state_key] = len(payload.get("users", [])) if isinstance(payload, dict) else 0
    elif state_key == "review_state" and isinstance(payload, dict):
      state_counts[state_key] = {
        "programmes": len(payload.get("programmeStatuses", {})),
        "edits": len(payload.get("programmeEdits", {})),
        "gaps": len(payload.get("gapStatuses", {})),
      }
    elif state_key == INSTITUTION_PROPOSAL_STATE_KEY and isinstance(payload, dict):
      proposals = payload.get("proposals", [])
      state_counts[state_key] = len(proposals) if isinstance(proposals, list) else 0
    else:
      state_counts[state_key] = 1
  return {
    "ok": True,
    "timestamp": now_iso(),
    "database": get_data_backend(),
    "database_ready": check_data_backend_ready(),
    "supabase_configured": supabase_configured(),
    "storage_ready": check_document_storage_ready(),
    "email_configured": smtp_configured(),
    "email_transport": email_transport(),
    "email_missing_keys": smtp_missing_keys(),
    "email_debug_codes": email_debug_codes_enabled(),
    "ai_configured": bool(get_gemini_api_key() if ai_provider == "gemini" else get_openai_api_key()),
    "ai_provider": ai_provider,
    "ai_model": ai_model,
    "storage_bucket": get_supabase_storage_bucket() if using_supabase() else None,
    "startup_persistence_error": STARTUP_PERSISTENCE_ERROR,
    "state_keys": [row.get("state_key") for row in state_rows],
    "state_counts": state_counts,
    "document_count": document_count,
    "documents_by_user": document_users,
    "document_extraction_statuses": extraction_statuses,
    "recommendation_run_count": run_count,
    "ai_chat_message_count": chat_message_count,
    "runtime_event_count": event_count,
    "latest_events": safe_list_runtime_events(5),
    "latest_runs": latest_runs,
  }


@app.post("/api/events")
def record_runtime_event(payload: RuntimeEventRequest, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  user = require_current_user(authorization)
  check_rate_limit(request, "runtime_event", 360, 3600, user["id"])
  event_type = payload.event_type or payload.eventType
  try:
    event = insert_runtime_event(user["id"], event_type or "", payload.label, payload.payload)
  except ValueError as exc:
    raise HTTPException(status_code=400, detail=str(exc))
  return {"ok": True, "eventId": event["id"]}


def call_gemini(payload: GuidanceRequest, request_payload: dict[str, Any]) -> dict[str, Any]:
  api_key = get_gemini_api_key()
  model = get_gemini_model_candidates()[0]
  if not api_key:
    return {
      "mode": "local_fallback",
      "provider": "gemini",
      "model": model,
      "guidance": build_fallback(payload, "GEMINI_API_KEY is not set, so local guidance was used."),
    }

  try:
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=api_key)
    last_error: Exception | None = None
    for candidate_model in get_gemini_model_candidates():
      try:
        response = client.models.generate_content(
          model=candidate_model,
          contents=json.dumps(request_payload, ensure_ascii=False),
          config=types.GenerateContentConfig(
            system_instruction=DEVELOPER_PROMPT,
            response_mime_type="application/json",
          ),
        )
        text = response.text or ""
        guidance = normalize_ai_guidance(
          parse_json_response(text) or build_fallback(payload, "The Gemini response was not valid JSON."),
          payload,
        )
        return {
          "mode": "gemini",
          "provider": "gemini",
          "model": candidate_model,
          "guidance": guidance,
        }
      except Exception as exc:
        last_error = exc
        if "NOT_FOUND" not in str(exc) and "not found" not in str(exc).lower() and "no longer available" not in str(exc).lower():
          raise
    raise RuntimeError(f"Gemini failed for available model fallbacks: {last_error}")
  except Exception as exc:
    return {
      "mode": "local_fallback",
      "provider": "gemini",
      "model": model,
      "guidance": build_fallback(payload, str(exc)),
    }


def call_openai(payload: GuidanceRequest, request_payload: dict[str, Any]) -> dict[str, Any]:
  api_key = get_openai_api_key()
  model = os.getenv("OPENAI_MODEL", "gpt-4o").strip() or "gpt-4o"
  if not api_key:
    return {
      "mode": "local_fallback",
      "provider": "openai",
      "model": model,
      "guidance": build_fallback(payload, "OPENAI_API_KEY is not set, so local guidance was used."),
    }

  try:
    client = OpenAI(api_key=api_key)
    response = client.responses.create(
      model=model,
      input=[
        {"role": "developer", "content": DEVELOPER_PROMPT},
        {"role": "user", "content": json.dumps(request_payload, ensure_ascii=False)},
      ],
    )
    text = response.output_text
    guidance = normalize_ai_guidance(
      parse_json_response(text) or build_fallback(payload, "The OpenAI response was not valid JSON."),
      payload,
    )
    return {
      "mode": "openai",
      "provider": "openai",
      "model": model,
      "guidance": guidance,
      "response_id": response.id,
    }
  except Exception as exc:
    return {
      "mode": "local_fallback",
      "provider": "openai",
      "model": model,
      "guidance": build_fallback(payload, str(exc)),
    }


@app.get("/api/ai/chat")
def ai_chat_history(authorization: str | None = Header(default=None)) -> dict[str, Any]:
  user = require_current_user(authorization)
  return {"ok": True, "messages": safe_list_ai_chat_history(user["id"])}


@app.delete("/api/ai/chat")
def ai_chat_clear(authorization: str | None = Header(default=None)) -> dict[str, Any]:
  user = require_current_user(authorization)
  safe_clear_ai_chat_history(user["id"])
  users = get_auth_users_internal()
  stored = next((item for item in users if item["id"] == user["id"]), None)
  if stored:
    add_user_activity(stored, "ai_chat_cleared", "Cleared EduGuide AI chat", user)
    save_auth_users_internal(users)
  return {"ok": True, "messages": []}


@app.post("/api/ai/guidance")
def ai_guidance(payload: GuidanceRequest, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
  current_user = require_current_user(authorization)
  safe_payload = validate_and_sanitize_guidance_payload(payload)
  if not safe_payload.matches and not safe_payload.blockedMatches:
    return {
      "mode": "local_fallback",
      "model": None,
      "guidance": build_fallback(safe_payload, "No matches were sent to the AI endpoint."),
    }
  check_rate_limit(request, "ai_guidance", 40, 3600, current_user["id"])
  incoming_history = merge_ai_chat_histories(safe_payload.conversation)
  server_history = safe_list_ai_chat_history(current_user["id"])
  merged_history = merge_ai_chat_histories(server_history, incoming_history)
  request_payload = build_request_payload(safe_payload, merged_history)
  save_recommendation_run(safe_payload, request_payload)
  if get_ai_provider() == "openai":
    result = call_openai(safe_payload, request_payload)
  else:
    result = call_gemini(safe_payload, request_payload)

  guidance = result.get("guidance") if isinstance(result, dict) else {}
  assistant_message = {
    "id": f"ai-{uuid.uuid4().hex[:12]}",
    "role": "assistant",
    "content": guidance_to_chat_text(guidance if isinstance(guidance, dict) else {}),
    "at": now_iso(),
  }
  chat = safe_save_ai_chat_history(current_user["id"], [*merged_history, assistant_message])
  users = get_auth_users_internal()
  stored = next((item for item in users if item["id"] == current_user["id"]), None)
  if stored:
    add_user_activity(stored, "ai_guidance", "Used EduGuide AI", current_user, {"mode": safe_payload.mode})
    save_auth_users_internal(users)
  result["chat"] = chat
  return result


@app.get("/health")
def health() -> dict[str, Any]:
  """Return only non-sensitive readiness flags for public uptime checks."""
  return {
    "ok": True,
    "database_ready": check_data_backend_ready(),
    "storage_ready": check_document_storage_ready(),
  }


@app.get("/")
def index() -> FileResponse:
  return cached_file_response(ROOT / "index.html", media_type="text/html", headers=NO_CACHE_HEADERS)


@app.get("/index.html")
def index_html() -> FileResponse:
  return cached_file_response(ROOT / "index.html", media_type="text/html", headers=NO_CACHE_HEADERS)


@app.get("/styles.css")
def styles() -> FileResponse:
  return cached_file_response(ROOT / "styles.css", media_type="text/css", headers=IMMUTABLE_CACHE_HEADERS)


@app.get("/app.js")
def app_script() -> FileResponse:
  return cached_file_response(ROOT / "app.js", media_type="application/javascript", headers=IMMUTABLE_CACHE_HEADERS)


@app.get("/manifest.webmanifest")
def web_manifest() -> FileResponse:
  return cached_file_response(ROOT / "manifest.webmanifest", media_type="application/manifest+json")


@app.get("/sw.js")
def service_worker() -> FileResponse:
  return cached_file_response(ROOT / "sw.js", media_type="application/javascript", headers=NO_CACHE_HEADERS)


@app.get("/icons/{file_name}")
def public_icon(file_name: str) -> FileResponse:
  if file_name not in PUBLIC_ICON_FILES:
    raise HTTPException(status_code=404, detail="Icon is not public.")
  path = ROOT / "icons" / file_name
  if not path.exists():
    raise HTTPException(status_code=404, detail="Icon not found.")
  return cached_file_response(path, media_type=guess_mime_type(path), headers=IMMUTABLE_CACHE_HEADERS)


@app.get("/data/{file_name}")
def public_data_file(file_name: str) -> FileResponse:
  if file_name not in PUBLIC_DATA_FILES:
    raise HTTPException(status_code=404, detail="File is not public.")
  path = ROOT / "data" / file_name
  if not path.exists():
    raise HTTPException(status_code=404, detail="File not found.")
  media_type = "application/javascript" if file_name.endswith(".js") else "application/json"
  return cached_file_response(path, media_type=media_type)
